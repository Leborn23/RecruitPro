from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
OUT_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATE_DIR = Path(r"D:/Program Files (x86)/jilu/xwechat_files/wxid_nztvtpqbm8qt22_9e13/msg/file/2026-04")

PROJECT_NAME = "RecruitPro智能招聘管理系统"
DOC_DATE = date(2026, 4, 26).strftime("%Y年%m月%d日")


def find_template(keyword: str) -> Path:
    for path in TEMPLATE_DIR.glob("*.docx"):
        if keyword in path.name:
            return path
    raise FileNotFoundError(f"未找到包含 {keyword} 的模板文件")


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def new_doc_from_template(keyword: str) -> Document:
    doc = Document(str(find_template(keyword)))
    clear_body(doc)
    compact_template_styles(doc)
    return doc


def compact_template_styles(doc: Document) -> None:
    for style_name in ["Normal", "List Paragraph"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.12
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.paragraph_format.page_break_before = False
            style.paragraph_format.keep_with_next = False
            style.paragraph_format.keep_together = False
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cm_to_twips(value: float) -> int:
    return int(value / 2.54 * 1440)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(cm_to_twips(width_cm)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(cm_to_twips(width_cm)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_cm: list[float]) -> None:
    tbl = table._tbl
    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(cm_to_twips(width)))
        grid.append(col)
    tbl.insert(1, grid)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(9)


def style_table(table, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed_layout(table)
    table.style = "Table Grid"
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "微软雅黑"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    r.font.size = Pt(9)
            if header and ri == 0:
                set_cell_shading(cell, "D9EAF7")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def apply_column_widths(table, widths_cm: list[float]) -> None:
    set_table_width(table, sum(widths_cm))
    set_table_grid(table, widths_cm)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 16, "16355F"),
        ("Heading 2", 13, "1F5FBF"),
        ("Heading 3", 11.5, "24476B"),
        ("Heading 4", 10.5, "355B87"),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    run = p.add_run(PROJECT_NAME)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("16355F")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(22)
    run = p2.add_run(title)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("1F5FBF")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(18)
    r = p3.add_run(subtitle)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("56718F")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(160)
    r = p4.add_run(DOC_DATE)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(12)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_template_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(PROJECT_NAME)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(105)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(22)

    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(22)

    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(DOC_DATE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    p.runs[0].font.size = Pt(12)
    doc.add_page_break()


def add_revision_table(doc: Document) -> None:
    doc.add_heading("修订记录", level=1)
    for key, value in [
        ("时间", DOC_DATE),
        ("版本", "V1.0"),
        ("作者", "Codex"),
        ("主要修订内容", "按模板章节结构重新生成，补充 RecruitPro 项目需求、概要设计、Agent、接口、数据、验收和交付内容。"),
    ]:
        para = doc.add_paragraph()
        para.add_run(f"{key}：").bold = True
        para.add_run(value)
    doc.add_paragraph()


def add_toc(doc: Document, rows: list[str]) -> None:
    p = doc.add_paragraph("目 录")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    for row in rows:
        if row.startswith("第"):
            style_name = "toc 1"
            indent = 0
        elif re.match(r"^\d+\.\d+\.\d+", row):
            style_name = "toc 3"
            indent = 0.9
        elif re.match(r"^\d+\.\d+", row):
            style_name = "toc 2"
            indent = 0.45
        else:
            style_name = "toc 2"
            indent = 0.2
        text = row if "\t" in row else f"{row}\t"
        p = doc.add_paragraph(text)
        if style_name in doc.styles:
            p.style = doc.styles[style_name]
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.02
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        for run in p.runs:
            run.font.size = Pt(10.2)
    doc.add_page_break()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    for k, v in rows:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run(f"{k}：").bold = True
        para.add_run(v)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.add_run("• ").bold = True
        p.add_run(item)


def normalize_body_alignment(doc: Document) -> None:
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.keep_together = False
        p.paragraph_format.page_break_before = False
        if p.alignment is None and style_name in {"Normal", "List Paragraph"}:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_image_if_exists(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(4.85))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"]


MODULES = [
    {
        "name": "登录与密码重置",
        "priority": "高",
        "desc": "支持用户通过 Supabase Auth 登录系统，并提供忘记密码、邮件回调和重置密码流程。",
        "pre": "用户访问 /login 或通过重置链接进入 /reset-password。",
        "post": "登录成功进入招聘指挥台；密码重置成功后可重新登录。",
        "ui": ROOT / "ui-screenshots" / "01-login.png",
    },
    {
        "name": "招聘指挥台",
        "priority": "高",
        "desc": "展示岗位总数、进行中岗位、候选人样本、近期面试，以及岗位、候选人和面试的快速入口。",
        "pre": "用户具备 VIEW_DASHBOARD 权限并完成登录。",
        "post": "用户可进入岗位、候选人或面试模块继续处理。",
        "ui": ROOT / "ui-screenshots" / "03-dashboard.png",
    },
    {
        "name": "岗位管理",
        "priority": "高",
        "desc": "支持维护岗位名称、部门、地点、状态、筛选阈值、技术要求、年龄、学历和经验要求。",
        "pre": "用户具备 MANAGE_POSITIONS 权限。",
        "post": "岗位数据写入 active_positions，并作为简历筛选、候选人匹配和面试安排的基础。",
        "ui": ROOT / "ui-screenshots" / "15-positions.png",
    },
    {
        "name": "简历筛选",
        "priority": "高",
        "desc": "支持批量上传 PDF/DOCX 简历，执行文本提取、结构化解析、岗位匹配、匹配分数计算和人工复核。",
        "pre": "用户选择目标岗位并上传符合格式的简历文件。",
        "post": "生成 resume_uploads、parsed_resume_profiles、parsed_resume_projects、candidate_position_matches 和 candidates 记录。",
        "ui": ROOT / "ui-screenshots" / "16-screening.png",
    },
    {
        "name": "候选人管理与详情",
        "priority": "高",
        "desc": "支持候选人列表、详情查看、匹配分解、证据片段、项目经历、风险提示、人工通过/待定/淘汰和删除。",
        "pre": "候选人由简历筛选流程生成，或已有候选人数据。",
        "post": "招聘人员可据此发起面试邀请或做人工复核结论。",
        "ui": ROOT / "ui-screenshots" / "25-candidate-detail.png",
    },
    {
        "name": "面试管理与 AI 面试",
        "priority": "高",
        "desc": "支持面试排期、面试房间、口令校验、问答轮次、AI Agent 自动追问、AI 评分、人工确认和面试报告查看。",
        "pre": "候选人已进入可面试状态，系统配置 FastAPI 后端和外部 Agent 服务。",
        "post": "生成 interview_sessions、interview_turns 和 interview_reports，并更新面试状态。",
        "ui": ROOT / "ui-screenshots" / "18-interviews.png",
    },
    {
        "name": "AI 面试 Agent",
        "priority": "高",
        "desc": "Agent 根据候选人简历、岗位 JD、结构化岗位要求和面试配置生成问题计划，负责首问、追问、状态推进、结束判断和最终报告生成。",
        "pre": "FastAPI 已配置 AGENT_BASE_URL、AGENT_SHARED_SECRET、AGENT_TIMEOUT_MS；面试会话已创建并具备 candidate_id、position_id、session_id。",
        "post": "Agent 返回 interview_plan、message、state_snapshot、final_report；系统将问题计划、AI 轮次、状态快照和评分报告写入数据库。",
        "ui": ROOT / "ui-screenshots" / "26-interview-room.png",
    },
    {
        "name": "系统设置与权限管理",
        "priority": "高",
        "desc": "支持组织设置、AI 策略、模型配置、面试题量/时长、管理员资料、角色和权限管理。",
        "pre": "用户具备 MANAGE_SETTINGS 权限；访问权限页需要超级管理员。",
        "post": "配置写入 company_settings、llm_model_configs、user_roles 等表并影响运行策略。",
        "ui": ROOT / "ui-screenshots" / "20-settings.png",
    },
    {
        "name": "薪酬市场与决策支持",
        "priority": "中",
        "desc": "支持导入本地薪酬市场文件，生成标准化记录、薪酬基准和候选人薪酬画像，为招聘决策提供参考。",
        "pre": "管理员准备 CSV/JSON/JSONL 薪酬数据，并配置 Supabase 服务密钥。",
        "post": "数据进入 market_salary_raw_records、market_salary_normalized_records 和 market_salary_benchmarks。",
        "ui": ROOT / "ui-screenshots" / "19-salary.png",
    },
]


def add_module_requirements(doc: Document) -> None:
    for idx, module in enumerate(MODULES, start=1):
        doc.add_heading(f"3.1.{idx} {module['name']}", level=3)
        doc.add_heading("需求说明", level=4)
        add_kv_table(
            doc,
            [
                ("模块描述", module["name"]),
                ("功能描述", module["desc"]),
                ("优先级", module["priority"]),
                ("输入/前置条件", module["pre"]),
                ("输出/后置条件", module["post"]),
                ("异常处理", "对权限不足、参数缺失、网络异常、文件超限、模型服务失败等情况给出明确提示，并保留必要状态便于重试或排查。"),
                ("业务规则", "所有关键业务数据需受登录态、权限键和数据库 RLS 共同约束；AI 输出必须保留证据、分数或人工复核入口。"),
            ],
        )
        doc.add_heading("页面图例", level=4)
        add_image_if_exists(doc, module["ui"], f"图 3.1.{idx}-1 {module['name']}页面")


def add_agent_requirements(doc: Document) -> None:
    doc.add_heading("3.1.10 AI 面试 Agent 专项需求", level=3)
    doc.add_heading("业务定位", level=4)
    doc.add_paragraph(
        "AI 面试 Agent 是 RecruitPro 面试流程中的智能问答与评估组件。它不直接操作业务数据库，而是通过 FastAPI 网关接收候选人上下文、岗位上下文和会话标识，返回问题计划、下一轮提示、状态快照和最终评估报告。FastAPI 负责鉴权、数据组装、调用 Agent、结果映射和持久化。"
    )
    doc.add_heading("功能需求", level=4)
    add_bullets(
        doc,
        [
            "面试启动时，Agent 应接收 session_id、resume_text、jd_text、question_count 等参数，并生成 interview_plan。",
            "Agent 生成的问题计划应包含题目文本、主题、期望关键点、答题提示或评价关注点，便于前端展示和后端追踪。",
            "候选人每提交一轮回答，FastAPI 将 user_answer 和 session_id 转发给 Agent，Agent 返回下一条 message、当前状态和 state_snapshot。",
            "Agent 应区分继续提问、追问、结束面试、等待人工复核等状态，系统根据状态写入 interview_turns 的 metadata。",
            "Agent 完成后应提供 final_report，包含 overall_score、hire_recommendation、strengths、concerns 或 detailed_evaluations 等结构化字段。",
            "系统需将 Agent 报告映射为 interview_reports，包含总分、推荐结论、风险分、维度分、优势、风险、证据和摘要。",
            "当 Agent 尚未产出 final_report 时，系统需生成待复核报告，标记 pending_human_review，避免面试结果丢失。",
            "人工确认环节可覆盖最终建议、补充备注，并写入 human_confirmed、human_confirmed_by、human_confirmed_at。",
        ],
    )


def add_requirements_deep_expansion(doc: Document) -> None:
    doc.add_heading("3.1.11 功能需求明细补充", level=3)
    for idx, module in enumerate(MODULES, start=1):
        doc.add_heading(f"3.1.11.{idx} {module['name']}明细", level=4)
        doc.add_paragraph(
            f"{module['name']}模块需要围绕招聘业务闭环提供稳定、可追踪、可复核的能力。模块优先级为{module['priority']}，主要面向具备相应权限的招聘人员、面试官或管理员使用。"
        )
        add_kv_table(
            doc,
            [
                ("业务流程", f"用户进入{module['name']}入口后，系统首先校验登录态和权限，再加载相关基础数据；用户提交业务动作后，系统完成参数校验、业务规则判断、数据库写入或后端流程触发，并将结果反馈到页面。"),
                ("核心数据", "涉及用户身份、岗位、候选人、简历上传、匹配结果、面试记录、Agent 状态、系统配置或薪酬数据等对象，具体字段以 Supabase 迁移和后端模型为准。"),
                ("权限控制", "前端通过 ProtectedRoute 控制页面入口，后端通过 require_user、权限键和数据库 RLS 控制敏感操作；高权限动作需要管理员或超级管理员身份。"),
                ("异常处理", "对登录失效、权限不足、字段缺失、关联对象不存在、数据库写入失败、外部服务超时和网络错误给出明确提示，关键长流程需保留可重试状态。"),
                ("验收标准", "用户能够完成该模块的主要业务动作；页面状态与数据库记录一致；异常场景有可理解提示；关键操作有状态或结果记录；权限不足用户无法越权访问。"),
            ],
        )

    doc.add_heading("3.1.12 招聘业务流程需求", level=3)
    business_flows = [
        ("岗位创建流程", "招聘负责人创建岗位，录入岗位名称、部门、地点、状态、筛选阈值、技术要求、年龄、学历和经验条件；系统保存岗位后，简历筛选模块可选择该岗位作为目标岗位。"),
        ("简历批量筛选流程", "招聘专员选择目标岗位后批量上传简历，系统以并发控制方式执行上传、文本解析、结构化抽取、岗位匹配、候选人入库和结果展示。"),
        ("人工复核流程", "系统生成匹配分和推荐结论后，招聘人员在候选人详情页查看匹配技能、缺失技能、证据片段、风险点和项目经历，并标记人工通过、待定或淘汰。"),
        ("面试邀约流程", "招聘人员从候选人详情或面试模块创建面试，设置面试阶段、时间、面试官和房间信息，必要时生成房间口令。"),
        ("AI 面试流程", "候选人进入面试房间后，系统创建会话并启动 Agent，Agent 根据简历和岗位上下文生成题目、追问并最终形成评分报告。"),
        ("报告确认流程", "面试完成后，招聘人员查看 Agent 评分报告，对建议结论进行人工确认或修正，最终结论进入后续招聘决策。"),
        ("薪酬参考流程", "管理员导入市场薪酬数据，系统标准化并生成薪酬基准，招聘人员结合候选人薪酬画像辅助 offer 决策。"),
        ("权限维护流程", "超级管理员维护角色和权限，系统即时影响页面可见性、业务操作权限和数据库访问范围。"),
    ]
    for name, desc in business_flows:
        doc.add_heading(name, level=4)
        doc.add_paragraph(desc)
        doc.add_paragraph("流程要求：每一步都需要有明确入口、状态反馈、失败提示和可追踪的数据记录；长流程不得只依赖前端内存状态。")

    doc.add_heading("3.1.13 Agent 详细业务规则", level=3)
    add_bullets(
        doc,
        [
            "Agent 问题生成必须以岗位 JD 和候选人简历为主要上下文，避免脱离目标岗位泛泛提问。",
            "Agent 需要根据 question_count 控制题量，不应无限追问；追问应服务于补充证据而不是重复提问。",
            "Agent 的每一轮输出都必须通过 FastAPI 记录到 interview_turns，便于复盘面试过程。",
            "Agent 生成的评分报告不能直接作为最终录用结论，系统必须提供人工确认入口。",
            "Agent 报告中的优势、风险、证据和建议需要结构化保存，不能仅保存一段自然语言总结。",
            "Agent 服务异常时不得删除已有会话和回答，应允许重新评分或人工补录结论。",
            "同一个 session_id 对应一个连续面试线程，禁止将多个候选人的会话混用。",
            "Agent 返回内容中如缺少 final_report，系统应生成待复核报告并提示人工处理。",
        ],
    )

    doc.add_heading("3.6 非功能需求补充", level=2)
    non_functionals = [
        ("安全性", "系统需要保护候选人简历、面试记录、薪酬信息和账号权限数据。服务密钥仅允许后端和离线脚本使用；普通浏览器端只能使用匿名密钥和受 RLS 约束的访问方式。"),
        ("可靠性", "简历筛选、Agent 面试和薪酬导入均属于长流程，必须记录阶段状态、错误码、错误消息和重试信息，避免用户刷新页面后流程不可追踪。"),
        ("可维护性", "前端页面、业务后端、数据库迁移和运行文档需要保持一致；新增字段应通过迁移管理，新增接口应同步后端模型和前端调用代码。"),
        ("可扩展性", "系统需支持后续接入更多模型、OCR 服务、Agent 策略、薪酬数据来源和招聘业务阶段，接口和数据结构应避免硬编码单一供应商。"),
        ("可用性", "页面需要提供清晰的空状态、加载状态、错误提示和成功反馈；批量处理任务需要展示队列、进度和失败原因。"),
        ("审计性", "权限变更、候选人人工复核、面试报告确认、模型调用和关键导入任务应保留操作者、时间和上下文信息。"),
        ("隐私性", "简历原文、候选人联系方式、面试回答和薪酬信息均属于敏感信息，应减少不必要的前端暴露和日志输出。"),
        ("兼容性", "前端应运行于现代浏览器；后端可运行于本地 Python 环境或容器环境；数据库以 Supabase PostgreSQL 为目标。"),
    ]
    for title, body in non_functionals:
        doc.add_heading(title, level=3)
        doc.add_paragraph(body)

    doc.add_heading("3.7 接口需求清单", level=2)
    api_requirements = [
        ("健康检查", "GET /api/health 用于验证后端服务可用。"),
        ("公司设置", "GET /api/settings/company、PATCH /api/settings/company 用于读取和更新公司级配置。"),
        ("管理员", "GET /api/admin/users、POST /api/admin/users/{target_user_id}/permissions、POST /api/admin/users/{target_user_id}/role 用于用户权限管理。"),
        ("岗位", "GET/POST/PATCH/DELETE /api/positions 用于岗位列表、创建、编辑和删除。"),
        ("筛选", "POST /api/screening/phase1、POST /api/screening/persist-phase1、POST /api/screening/rescreen 用于简历处理和历史重筛。"),
        ("候选人", "GET /api/candidates、GET /api/candidates/{candidate_id}/detail、DELETE /api/candidates/{candidate_id} 用于候选人查询和管理。"),
        ("上传", "POST /api/uploads、PATCH /api/uploads/{upload_id}、POST /api/uploads/delete、POST /api/uploads/{upload_id}/cancel 用于上传状态管理。"),
        ("面试基础", "GET/POST/PATCH/DELETE /api/interviews 用于面试安排管理。"),
        ("Agent 面试", "POST /api/interviews/prepare、start、turn、finish、score、human-confirm 用于 AI 面试完整运行链路。"),
        ("房间口令", "POST /api/interviews/room-password 用于签发或校验面试房间口令。"),
        ("薪酬", "POST /api/salary/market/import、POST /api/salary/market/refresh、GET /api/salary/decision-dashboard 用于薪酬市场导入和决策展示。"),
    ]
    for name, desc in api_requirements:
        doc.add_paragraph(f"{name}：{desc}")
    doc.add_heading("输入输出", level=4)
    add_kv_table(
        doc,
        [
            ("输入数据", "候选人基本信息、结构化简历、项目经历、岗位信息、结构化 JD、面试题量、session_id 和候选人回答。"),
            ("输出数据", "问题计划、AI 话术、追问或结束提示、状态快照、最终评分报告、推荐结论、风险项和证据。"),
            ("状态数据", "asked_question_count、answer_count、next_nodes、agent_status、question_id、source=agent 等元数据。"),
            ("落库位置", "interview_sessions.question_plan、interview_turns.metadata、interview_reports.dimension_scores/evidence/summary 等字段。"),
        ],
    )
    doc.add_heading("异常与降级", level=4)
    add_bullets(
        doc,
        [
            "Agent 服务不可用、超时或返回非 JSON 时，FastAPI 应返回明确错误，并将会话保持在可恢复状态。",
            "Agent 返回 4xx/5xx 时，FastAPI 透出 detail 或 response 文本，便于前端提示和后端排查。",
            "Agent 已完成问答但报告未生成时，评分接口应创建待复核报告，供人工继续处理。",
            "共享密钥缺失或错误时，应禁止 Agent 调用成功，避免未授权服务接入。",
        ],
    )


def add_requirements_full_appendix(doc: Document) -> None:
    doc.add_heading("第五章 详细功能规格", level=1)
    detail_templates = [
        ("页面与入口", "系统需要在导航、按钮、列表操作和详情入口中体现该模块能力。入口展示必须与权限保持一致，用户无权限时不应看到可执行操作；如通过地址栏直接访问，系统应拦截并给出权限提示。"),
        ("数据加载", "进入页面后应加载当前用户可见的数据范围。列表数据需要支持空状态、加载状态和失败状态；详情数据需要在关联对象不存在、已删除或权限不足时给出明确反馈。"),
        ("业务操作", "核心操作需要包含参数校验、二次确认或明确反馈。对可能造成业务影响的操作，如删除、覆盖、批量处理、人工结论修改，需要保证操作结果可追踪。"),
        ("状态反馈", "页面需要展示处理进度、成功提示、失败原因和可继续操作。对于长流程，不能只显示旋转加载，需要将后端阶段状态转换为用户可理解的文案。"),
        ("权限与审计", "模块行为应遵守角色层级和权限键约束。敏感操作应记录操作者、操作时间、目标对象和关键结果，便于后续审计和问题追溯。"),
        ("异常场景", "应覆盖登录过期、权限不足、网络超时、数据库写入失败、外部服务异常、数据结构缺失和重复提交等场景，保证异常后页面仍可继续使用。"),
        ("验收要点", "验收时需要分别验证正常流程、空数据流程、失败流程、权限不足流程和刷新恢复流程；验收数据应覆盖至少一个真实岗位、多个候选人和一次完整面试链路。"),
    ]
    for idx, module in enumerate(MODULES, start=1):
        doc.add_heading(f"5.{idx} {module['name']}详细规格", level=2)
        doc.add_paragraph(
            f"{module['name']}属于{PROJECT_NAME}核心招聘流程的一部分，业务描述为：{module['desc']}。本节从页面、数据、操作、权限、异常和验收角度细化该模块要求。"
        )
        for title, body in detail_templates:
            doc.add_heading(title, level=3)
            doc.add_paragraph(body)
        doc.add_heading("模块专属补充", level=3)
        if "Agent" in module["name"]:
            add_bullets(
                doc,
                [
                    "启动面试前必须完成候选人、岗位、面试安排和会话记录校验，避免 Agent 收到不完整上下文。",
                    "每一轮候选人回答和 Agent 回复都必须持久化，页面刷新后能够恢复已发生的问答过程。",
                    "Agent 输出的题目、追问和结束语应在 metadata 中记录类型，供报告和复盘使用。",
                    "最终报告需要支持人工确认，人工确认结果优先作为招聘决策依据。",
                ],
            )
        elif "简历" in module["name"]:
            add_bullets(
                doc,
                [
                    "批量上传需要限制文件类型和大小，重复文件应通过 hash 或文件信息辅助识别。",
                    "处理链路至少覆盖上传、文本提取、结构化解析、岗位匹配、入库和完成阶段。",
                    "失败记录需要保留错误码、错误消息和可重试标识，便于招聘人员重新处理。",
                    "匹配结果需要展示分数来源、命中技能、缺失技能和证据片段。",
                ],
            )
        elif "权限" in module["name"]:
            add_bullets(
                doc,
                [
                    "首次超级管理员认领需要受到系统状态约束，已有超级管理员后不得重复认领。",
                    "角色修改和权限修改需要通过后端受控接口或 RPC 完成，避免普通客户端直接写敏感表。",
                    "页面权限、接口权限和数据库 RLS 需要形成一致约束，不能只依赖前端隐藏按钮。",
                ],
            )
        else:
            add_bullets(
                doc,
                [
                    "模块数据应与招聘主流程保持关联，避免形成无法回溯的孤立记录。",
                    "列表、详情和操作结果需要使用同一数据来源或明确刷新策略，避免页面显示与数据库状态不一致。",
                    "涉及 AI 结果的字段必须保留人工复核入口，避免将模型输出直接视为最终结论。",
                ],
            )

    doc.add_heading("第六章 数据字段与数据质量需求", level=1)
    data_quality_rows = [
        ("active_positions", "岗位名称、部门、地点、状态、筛选阈值、技术要求、年龄、学历、经验要求、创建时间、更新时间", "岗位名称、状态和筛选阈值不能为空；筛选阈值需在合理范围内；停用岗位不得进入新的筛选任务。"),
        ("candidates", "姓名、当前职位、经验、学历、年龄、城市、标签、亮点、关联岗位、匹配分、来源上传记录", "候选人需能回溯到岗位或上传来源；匹配分来自 AI 或人工复核时需保留来源。"),
        ("resume_uploads", "文件名、路径、大小、hash、状态、pipeline_stage、错误码、错误信息、重试次数", "状态必须与阶段一致；失败记录需要可读错误；取消或失败后不得被误认为完成。"),
        ("parsed_resume_profiles", "候选人画像、教育经历、技能、工作经历、联系方式、置信度、原始模型输出", "结构化字段应允许部分缺失，但必须保留置信度和原始载荷便于复核。"),
        ("parsed_resume_projects", "项目名称、项目摘要、技术栈、复杂度、职责、起止时间、候选人关联", "项目记录必须关联候选人或简历画像，避免无法归属的项目数据。"),
        ("candidate_position_matches", "候选人、岗位、总分、推荐结论、技能匹配、缺失项、证据、风险、人工复核", "匹配结果必须同时保存机器结论和人工结论，人工结论不应覆盖原始 AI 证据。"),
        ("upcoming_interviews", "候选人、岗位、面试阶段、时间、面试官、房间类型、房间口令、状态、报告关联", "时间、候选人和阶段不能为空；已完成面试应关联报告或待复核状态。"),
        ("interview_sessions", "面试 ID、会话状态、题目计划、上下文载荷、开始时间、结束时间、运行模式", "一个面试应有可识别会话；题目计划来自 Agent 时需记录 source。"),
        ("interview_turns", "会话 ID、轮次号、说话人、内容、输入模式、延迟、token、置信度、metadata", "轮次号需按会话递增；AI 和候选人轮次都需完整保存。"),
        ("interview_reports", "总分、推荐结论、维度分、优势、风险、证据、摘要、人工确认、确认人、确认时间", "报告必须能解释结论来源；人工确认字段需要独立保存，不得丢失 AI 原始结论。"),
        ("market_salary_raw_records", "来源、职位、城市、薪酬文本、公司、采集时间、原始载荷", "原始记录不可在标准化失败时丢失，需保留失败原因以便修复映射规则。"),
        ("market_salary_benchmarks", "标准岗位、城市、样本量、分位数、更新时间、来源范围", "样本量过低时应标记置信风险，供 offer 决策时参考。"),
        ("user_roles", "用户 ID、邮箱、角色、权限键、创建时间、更新时间", "角色层级必须受控；权限键只能来自系统定义集合。"),
        ("company_settings", "AI 策略、默认模型、简历隐私、反馈要求、面试配置、题量配置", "设置变更应即时影响后续流程，历史记录不应被错误重算。"),
        ("llm_usage_events", "场景、模型、token、耗时、成功状态、错误信息、关联业务对象", "模型调用统计应避免记录完整敏感简历明文，保留必要关联和摘要即可。"),
    ]
    for table_name, fields, rule in data_quality_rows:
        doc.add_heading(f"6.x {table_name}", level=2)
        doc.add_paragraph(f"主要字段：{fields}。")
        doc.add_paragraph(f"数据质量要求：{rule}")

    doc.add_heading("第七章 接口与集成需求扩展", level=1)
    interface_groups = [
        ("前端到 FastAPI", ["所有业务接口应返回统一错误信息结构。", "需要登录态的接口必须校验用户身份。", "涉及管理员、权限和服务密钥能力的接口不得由前端直接绕过后端访问数据库。"]),
        ("FastAPI 到 Supabase", ["服务端使用 service role key 执行受控写入。", "写入前应校验业务对象存在性和当前用户权限。", "数据库异常需要转换为业务可理解错误，不能直接暴露敏感连接信息。"]),
        ("FastAPI 到 Agent", ["调用必须携带 session_id 和必要上下文。", "请求头应包含 x-agent-secret。", "超时、非 JSON、错误状态码都应被捕获并反馈到面试流程。"]),
        ("前端到 Supabase", ["仅允许使用匿名密钥和 RLS 保护下的数据访问。", "页面订阅或查询结果必须遵守权限边界。", "敏感写入优先通过后端接口完成。"]),
        ("导入脚本到数据库", ["薪酬导入和批量数据任务需要使用受控环境变量。", "导入任务应记录来源、条数、成功数、失败数和错误明细。", "重复导入需要具备幂等或可清理策略。"]),
    ]
    for group, rules in interface_groups:
        doc.add_heading(group, level=2)
        add_bullets(doc, rules)

    doc.add_heading("第八章 测试与验收需求", level=1)
    acceptance_cases = [
        ("登录与权限", "使用普通用户、管理员、超级管理员分别登录，检查页面可见性、操作按钮、接口返回和数据库访问是否符合权限设计。"),
        ("岗位管理", "创建、编辑、停用和删除岗位，检查岗位列表、详情、筛选模块下拉选择和关联候选人显示是否同步。"),
        ("简历上传", "上传单个有效简历、多个有效简历、重复简历、超大文件和不支持格式文件，检查状态、错误提示和可重试能力。"),
        ("简历筛选", "对同一岗位执行首次筛选和历史重筛，检查匹配分、推荐结论、证据片段和人工复核字段。"),
        ("候选人详情", "打开候选人详情，检查基础信息、项目经历、匹配证据、薪酬画像、面试历史和操作入口。"),
        ("面试安排", "创建面试、修改时间、取消面试、进入房间，检查状态流转和房间口令逻辑。"),
        ("AI 面试 Agent", "完成准备、启动、至少两轮问答、结束、评分和人工确认，检查 interview_sessions、interview_turns 和 interview_reports 写入。"),
        ("Agent 异常", "模拟 Agent 超时、返回错误、无 final_report 和密钥错误，检查系统提示、待复核报告和会话保留情况。"),
        ("薪酬市场", "导入薪酬原始数据、刷新标准化记录和基准，检查样本量、分位数和决策看板显示。"),
        ("系统设置", "修改公司设置、AI 策略和默认模型，检查后续筛选和面试流程是否读取最新配置。"),
    ]
    for idx, (name, desc) in enumerate(acceptance_cases, start=1):
        doc.add_heading(f"8.{idx} {name}验收", level=2)
        doc.add_paragraph(desc)
        add_bullets(
            doc,
            [
                "前置条件：准备具备代表性的测试账号、岗位、候选人、简历文件和必要环境变量。",
                "执行过程：按业务页面正常操作，同时观察后端日志、数据库记录和页面反馈。",
                "通过标准：功能结果正确、状态可追踪、异常提示明确、无越权访问、刷新后数据仍一致。",
            ],
        )

    doc.add_heading("第九章 运维与部署需求", level=1)
    ops_sections = [
        ("环境变量管理", "前端、后端、Agent、Supabase 和导入脚本的环境变量应按环境隔离。生产环境不得使用开发密钥；服务密钥不得写入前端构建产物。"),
        ("数据库迁移", "所有结构变更必须通过 supabase/migrations 管理。上线前需要确认迁移顺序、RLS 策略、RPC 函数和索引均已应用。"),
        ("日志与监控", "后端应记录接口耗时、Agent 调用耗时、模型调用结果、导入任务统计和关键错误。日志需避免输出完整简历明文和房间口令。"),
        ("备份与恢复", "数据库和对象存储需要定期备份。恢复演练需覆盖候选人、简历文件、面试报告和权限配置。"),
        ("容量规划", "随着简历文件、面试轮次、模型调用记录和薪酬原始数据增长，需要关注数据库表大小、索引效率和 Storage 容量。"),
        ("发布检查", "发布前需执行前端构建、后端健康检查、数据库迁移检查、Agent 连通性检查和至少一条端到端招聘流程验证。"),
    ]
    for title, body in ops_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)


def add_design_full_appendix(doc: Document) -> None:
    doc.add_heading("第八章 模块详细设计", level=1)
    for idx, module in enumerate(MODULES, start=1):
        doc.add_heading(f"8.{idx} {module['name']}设计", level=2)
        doc.add_paragraph(f"设计目标：{module['desc']}")
        design_points = [
            ("前端设计", "页面采用 React 组件组织，路由入口由权限保护组件控制。页面内部应拆分为列表区、详情区、操作区、状态提示区和空状态区，避免在单个组件中堆积过多业务分支。"),
            ("后端设计", "复杂业务、服务密钥操作、长流程编排和外部服务调用集中在 FastAPI 后端实现。后端接口负责参数校验、权限校验、业务规则执行、数据库写入和错误转换。"),
            ("数据设计", "模块数据应优先写入对应业务表，并通过外键或业务 ID 关联岗位、候选人、上传记录、面试会话或用户。AI 输出需要保留原始载荷摘要和结构化字段。"),
            ("状态设计", "模块状态应使用有限集合表达，避免用自由文本作为流程判断依据。前端只根据后端返回状态渲染，不应自行推断长流程结果。"),
            ("安全设计", "页面级权限、接口级权限和数据库级策略共同约束。敏感字段、服务密钥、管理员操作和候选人隐私信息不得在无权限场景暴露。"),
            ("可测试设计", "每个模块至少保留可构造的测试入口和稳定数据状态。长流程通过状态字段验证，AI 相关流程可通过模拟服务验证异常和降级。"),
        ]
        for title, body in design_points:
            doc.add_heading(title, level=3)
            doc.add_paragraph(body)
        if "Agent" in module["name"]:
            add_bullets(
                doc,
                [
                    "Agent 不直接连接业务数据库，所有业务持久化由 FastAPI 完成。",
                    "agent_fetch 统一封装请求地址、共享密钥、超时和 JSON 解析，避免分散调用逻辑。",
                    "Agent 返回的 state_snapshot 写入 interview_turns.metadata，报告结果通过映射函数写入 interview_reports。",
                ],
            )
        elif "薪酬" in module["name"]:
            add_bullets(
                doc,
                [
                    "薪酬原始记录、标准化记录和基准结果分层保存，避免标准化失败影响原始数据追溯。",
                    "导入任务记录成功数、失败数和错误明细，支持运营人员定位数据质量问题。",
                ],
            )
        else:
            add_bullets(
                doc,
                [
                    "模块实现需要与现有 Supabase 表结构和前端类型定义保持一致。",
                    "页面组件应复用现有布局、按钮、提示和图标风格，保证整体操作体验一致。",
                ],
            )

    doc.add_heading("第九章 Agent 详细设计", level=1)
    doc.add_heading("9.1 组件职责", level=2)
    add_kv_table(
        doc,
        [
            ("React 面试页面", "负责展示面试房间、题目、候选人回答输入、轮次列表、状态提示和报告入口。"),
            ("FastAPI 面试接口", "负责准备会话、组装上下文、调用 Agent、保存轮次、映射报告和人工确认。"),
            ("Agent 服务", "负责基于简历与 JD 生成问题、追问、状态快照和最终评估报告。"),
            ("Supabase 数据库", "负责保存面试安排、会话、轮次、报告、候选人和岗位数据。"),
        ],
    )
    doc.add_heading("9.2 状态机设计", level=2)
    agent_states = [
        ("preparing", "准备候选人和岗位上下文，创建或复用 interview_sessions。"),
        ("ready", "面试准备完成，等待候选人启动。"),
        ("in_progress", "Agent 已启动并正在进行多轮问答。"),
        ("waiting_agent", "候选人已提交回答，系统等待 Agent 返回下一步。"),
        ("scoring", "问答结束，系统请求或等待最终报告。"),
        ("pending_human_review", "Agent 未返回完整报告或需要人工复核。"),
        ("done", "报告已生成并完成必要确认。"),
        ("failed", "Agent、网络或数据异常导致流程无法自动继续。"),
    ]
    for state, desc in agent_states:
        doc.add_paragraph(f"{state}：{desc}")
    doc.add_heading("9.3 调用时序设计", level=2)
    sequence_steps = [
        "招聘人员创建面试安排，系统写入 upcoming_interviews。",
        "候选人或招聘人员进入面试房间，前端调用 /api/interviews/prepare。",
        "FastAPI 校验面试、候选人、岗位和权限，创建 interview_sessions。",
        "前端调用 /api/interviews/start，FastAPI 组装 resume_text 和 jd_text。",
        "FastAPI 调用 Agent /agent/start，携带 session_id、上下文和题量。",
        "Agent 返回 interview_plan 和第一条 message，FastAPI 写入 question_plan 和首轮 AI turn。",
        "候选人提交回答，前端调用 /api/interviews/turn。",
        "FastAPI 写入候选人 turn，调用 Agent /agent/answer，并写入下一条 AI turn。",
        "Agent 判断题量和追问需要，返回继续、追问、结束或待复核状态。",
        "面试结束后，前端调用 /api/interviews/score，FastAPI 查询 Agent /agent/status。",
        "FastAPI 将 final_report 映射为 interview_reports，或生成 pending_human_review 报告。",
        "招聘人员进行人工确认，系统写入 human_confirmed 相关字段。",
    ]
    for step_no, step in enumerate(sequence_steps, start=1):
        doc.add_paragraph(f"{step_no}. {step}")
    doc.add_heading("9.4 报告映射设计", level=2)
    mapping_rows = [
        ("overall_score", "interview_reports.overall_score", "总分映射，数值范围需归一到系统展示标准。"),
        ("hire_recommendation", "interview_reports.recommendation", "通过 map_agent_recommendation 归一为 hire、hold、needs_review、reject。"),
        ("strengths", "interview_reports.strengths", "保留候选人优势，展示给招聘人员复核。"),
        ("concerns/risks", "interview_reports.risks", "保留风险点并参与风险分计算。"),
        ("dimension_scores", "interview_reports.dimension_scores", "保存岗位匹配、技术深度、项目证据、沟通表达等维度。"),
        ("detailed_evaluations", "interview_reports.question_evaluations", "保存每道题的回答、反馈、证据和评分。"),
        ("state_snapshot", "interview_reports.evidence", "无 final_report 或需复核时作为证据载荷保存。"),
    ]
    for src, dst, rule in mapping_rows:
        doc.add_paragraph(f"{src} -> {dst}：{rule}")

    doc.add_heading("第十章 数据库详细设计", level=1)
    database_design = [
        ("岗位域", "active_positions 保存岗位主体；parsed_job_requirements 保存结构化 JD；二者通过岗位 ID 关联，支持岗位规则复核和历史重筛。"),
        ("候选人域", "candidates 保存候选人主档；parsed_resume_profiles 和 parsed_resume_projects 保存结构化简历和项目经历；candidate_position_matches 保存岗位匹配结果。"),
        ("上传域", "resume_uploads 是简历处理流水线入口，记录文件信息、处理阶段、错误信息和重试次数。"),
        ("面试域", "upcoming_interviews 保存安排；interview_sessions 保存会话；interview_turns 保存问答轮次；interview_reports 保存评分报告。"),
        ("薪酬域", "market_salary_raw_records、market_salary_normalized_records、market_salary_benchmarks 和 market_salary_crawl_jobs 分别承担原始、标准化、基准和任务记录职责。"),
        ("权限域", "user_roles 保存角色和权限键；公司设置、模型配置和模型调用记录分别由 company_settings、llm_model_configs 和 llm_usage_events 支撑。"),
    ]
    for title, body in database_design:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
        add_bullets(
            doc,
            [
                "表结构变更通过迁移文件控制，避免生产环境手工改表造成版本不一致。",
                "涉及用户和候选人的表需要启用或配合 RLS 策略，敏感写入通过后端完成。",
                "重要 JSON 字段用于保留 AI 载荷和证据，但核心查询字段应结构化保存以便检索和统计。",
            ],
        )

    doc.add_heading("第十一章 错误处理与日志设计", level=1)
    error_design = [
        ("输入错误", "字段缺失、格式错误、文件类型不支持、题量非法等由接口参数校验拦截，前端展示可修改提示。"),
        ("权限错误", "未登录、权限不足、角色层级不满足时返回 401 或 403，并在页面上隐藏不可操作入口。"),
        ("数据错误", "关联岗位、候选人、面试、上传记录不存在时返回明确错误，不允许创建孤立业务数据。"),
        ("外部服务错误", "LLM、OCR、Agent 超时或返回错误时，后端记录请求场景、耗时、状态码和错误摘要。"),
        ("数据库错误", "数据库写入失败、迁移缺失、约束冲突时，系统应保留原流程状态并返回可排查信息。"),
        ("前端错误", "前端捕获接口失败和渲染异常，展示错误状态并提供重试或返回入口。"),
    ]
    for title, body in error_design:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
    doc.add_heading("日志字段建议", level=2)
    add_bullets(
        doc,
        [
            "request_id、user_id、endpoint、method、status_code、duration_ms。",
            "business_object_type、business_object_id、position_id、candidate_id、interview_id、session_id。",
            "external_service、external_status、timeout_ms、retry_count、error_code、error_message。",
            "日志中不得记录完整服务密钥、房间口令、完整简历明文和不必要的候选人联系方式。",
        ],
    )

    doc.add_heading("第十二章 测试设计", level=1)
    test_designs = [
        ("单元测试", "后端映射函数、薪酬标准化函数、Agent 推荐结论归一化、权限判断和数据转换逻辑应优先做单元测试。"),
        ("接口测试", "覆盖健康检查、岗位、候选人、上传、面试、Agent、薪酬和管理员接口，验证正常、异常和权限不足场景。"),
        ("集成测试", "覆盖简历上传到候选人入库、AI 面试到报告生成、薪酬导入到基准刷新等跨模块流程。"),
        ("前端测试", "覆盖路由保护、列表空状态、表单校验、批量任务状态、面试房间交互和报告确认操作。"),
        ("安全测试", "验证匿名用户、普通用户、管理员和超级管理员在页面、接口和数据库访问上的边界。"),
        ("回归测试", "每次修改数据库迁移、Agent 接口、简历处理或权限逻辑后，必须回归招聘主流程。"),
    ]
    for title, body in test_designs:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
        doc.add_paragraph("测试输出应包括执行环境、测试账号、测试数据、执行结果、失败截图或日志摘要，以及是否影响上线的结论。")

    doc.add_heading("第十三章 接口详细设计补充", level=1)
    endpoint_designs = [
        ("GET /api/health", "用于服务存活检查。接口不依赖业务登录态，应返回服务状态、版本或基础可用信息，供部署检查和监控探测使用。"),
        ("GET /api/dashboard", "聚合岗位、候选人、面试和筛选概览。后端负责控制查询范围，前端只展示返回结果，不在浏览器端拼装敏感统计。"),
        ("GET/POST/PATCH/DELETE /api/positions", "岗位管理接口。创建和修改时校验岗位名称、状态、筛选阈值和要求字段；删除或停用时需考虑关联候选人和筛选历史。"),
        ("POST /api/screening/phase1", "启动第一阶段简历筛选。接口接收岗位和文件信息，创建处理任务，推进文本提取、结构化解析和匹配计算。"),
        ("POST /api/screening/persist-phase1", "将阶段一筛选结果固化为候选人和匹配记录。接口需保证幂等，避免重复写入同一候选人与同一岗位的重复结果。"),
        ("POST /api/screening/rescreen", "对历史候选人或上传记录按最新岗位规则重新筛选。接口需记录重筛来源，避免覆盖历史评估证据。"),
        ("GET /api/candidates", "候选人列表查询。支持按岗位、状态、标签、匹配分和关键词筛选，返回字段需控制敏感信息暴露范围。"),
        ("GET /api/candidates/{candidate_id}/detail", "候选人详情查询。聚合候选人主档、结构化简历、项目经历、岗位匹配、面试记录和薪酬画像。"),
        ("GET/POST/PATCH/DELETE /api/interviews", "面试安排管理。用于创建、调整、取消和查询面试，状态变更需与 AI 面试会话和报告关联保持一致。"),
        ("POST /api/interviews/prepare", "准备 AI 面试会话。接口校验候选人、岗位、面试记录和房间权限，创建或返回 interview_sessions。"),
        ("POST /api/interviews/start", "启动 Agent 面试。接口组装上下文并调用 /agent/start，保存题目计划和首轮 AI 话术。"),
        ("POST /api/interviews/turn", "提交候选人回答并请求 Agent 下一轮输出。接口先保存候选人回答，再调用 Agent，最后保存 AI 回复和状态快照。"),
        ("POST /api/interviews/score", "触发或获取 Agent 最终报告。接口读取 /agent/status，完成报告映射或生成待人工复核报告。"),
        ("POST /api/interviews/human-confirm", "人工确认面试报告。接口保存人工建议、备注、确认人和确认时间，并更新面试报告状态。"),
        ("POST /api/salary/market/import", "导入薪酬市场数据。接口或脚本需要将原始数据、标准化结果和导入任务分开记录。"),
        ("GET /api/salary/decision-dashboard", "返回薪酬决策看板数据。输出应包括岗位、城市、样本量、分位数和候选人薪酬参考。"),
    ]
    for endpoint, body in endpoint_designs:
        doc.add_heading(endpoint, level=2)
        doc.add_paragraph(body)
        add_bullets(
            doc,
            [
                "输入设计：使用明确字段名和类型，必填字段由后端统一校验。",
                "输出设计：成功时返回业务对象或状态摘要，失败时返回可读 detail。",
                "安全设计：涉及业务数据的接口必须校验登录态和权限，敏感写入只允许后端服务完成。",
            ],
        )

    doc.add_heading("第十四章 部署架构详细设计", level=1)
    deployment_parts = [
        ("前端静态站点", "前端通过 npm run build 生成 dist 目录。部署时需要注入 VITE_SUPABASE_URL、VITE_SUPABASE_ANON_KEY 和 VITE_API_BASE_URL。前端不包含服务密钥，也不直接保存 Agent 密钥。"),
        ("FastAPI 服务", "后端负责业务编排和敏感写入。部署时需要配置 SUPABASE_URL、SUPABASE_SERVICE_ROLE_KEY、AGENT_BASE_URL、AGENT_SHARED_SECRET 和 AGENT_TIMEOUT_MS。生产环境应通过进程管理器或容器运行。"),
        ("Agent 服务", "Agent 作为独立 HTTP 服务运行，只暴露给 FastAPI。它提供 /agent/start、/agent/answer 和 /agent/status，校验 x-agent-secret 并维护会话线程。"),
        ("Supabase 数据层", "Supabase 提供 PostgreSQL、Auth、Storage、RLS 和 RPC。上线前必须执行迁移，确认 Storage bucket、RLS 策略和权限函数可用。"),
        ("离线任务与脚本", "薪酬导入、数据修复和批量迁移脚本使用服务密钥运行。脚本应只在可信环境执行，并记录操作日志和结果文件。"),
    ]
    for title, body in deployment_parts:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
        add_bullets(
            doc,
            [
                "配置项应按开发、测试、生产环境隔离。",
                "启动顺序建议为数据库迁移、Agent 服务、FastAPI 服务、前端站点。",
                "发布后需执行健康检查、Agent 连通性检查和端到端招聘流程冒烟测试。",
            ],
        )
    doc.add_heading("网络与安全边界", level=2)
    add_bullets(
        doc,
        [
            "浏览器只能访问前端、FastAPI 公开业务接口和 Supabase 匿名密钥允许的资源。",
            "FastAPI 可以访问 Supabase 服务端接口、Agent 服务和必要的模型/OCR 服务。",
            "Agent 不应暴露到公网匿名访问范围内，生产环境建议通过内网、网关或防火墙限制来源。",
            "服务密钥、共享密钥和数据库连接信息应通过部署平台密钥管理能力注入，不写入源码。",
        ],
    )

    doc.add_heading("第十五章 数据迁移与版本管理设计", level=1)
    migration_designs = [
        ("迁移来源", "supabase/migrations 是数据库结构、策略、函数和索引的唯一版本化来源。任何字段、表、RLS 或 RPC 变更均应通过迁移提交。"),
        ("迁移顺序", "迁移文件按时间顺序执行。新增表应先创建基础结构，再添加索引、策略、函数和数据修复脚本，避免部署中间态不可用。"),
        ("兼容策略", "新增字段优先采用可空字段或默认值，前端和后端代码发布后再逐步收紧约束，降低灰度发布风险。"),
        ("数据修复", "历史数据修复需记录修复范围、执行人、执行时间和影响条数。涉及候选人、面试报告和薪酬数据时应先备份。"),
        ("回滚设计", "结构性回滚需要评估数据丢失风险。无法直接回滚的迁移应提供修复脚本或向前兼容方案。"),
        ("版本核对", "部署前后需核对前端类型、后端模型、数据库字段和文档说明是否一致，避免接口存在但字段缺失。"),
    ]
    for title, body in migration_designs:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
    doc.add_heading("版本发布检查清单", level=2)
    add_bullets(
        doc,
        [
            "前端构建通过，关键页面可访问。",
            "后端健康检查通过，管理员接口、岗位接口和面试接口可用。",
            "数据库迁移已应用，RLS 策略和 RPC 函数存在。",
            "Agent 共享密钥一致，start、answer、status 三类接口均能响应。",
            "简历上传、候选人详情、AI 面试、人工确认和薪酬看板完成冒烟验证。",
        ],
    )

    doc.add_heading("第十六章 可维护性与扩展设计", level=1)
    maintainability_sections = [
        ("前端扩展", "新增页面应沿用现有路由、权限保护和布局体系。新增业务状态应在类型定义中声明，避免在多个页面使用不同字符串。"),
        ("后端扩展", "新增业务接口应集中在 FastAPI 中完成鉴权、参数校验和错误转换。外部服务调用应通过统一封装，便于超时、重试和日志治理。"),
        ("Agent 扩展", "Agent 可扩展不同面试策略，如技术面、项目深挖、行为面和管理面。扩展时应保持 final_report 映射字段稳定。"),
        ("模型扩展", "系统可通过 llm_model_configs 支持多个模型供应商和模型版本。调用记录写入 llm_usage_events，支持成本和质量分析。"),
        ("薪酬数据扩展", "薪酬数据源可从文件导入扩展为爬取任务或第三方 API。原始层、标准化层和基准层的分层设计保持不变。"),
        ("权限扩展", "新增权限键应同时更新前端路由、后端校验和管理员配置页面，避免出现页面可见但接口不可用的状态。"),
    ]
    for title, body in maintainability_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
        doc.add_paragraph("设计要求：扩展时必须保持现有数据兼容，不破坏已完成招聘流程、历史面试报告和人工确认结果。")


def add_requirements_extra_detail(doc: Document) -> None:
    doc.add_heading("第十章 用户用例详细说明", level=1)
    use_cases = [
        ("UC-01 用户登录系统", "招聘人员使用邮箱和密码登录系统，系统通过 Supabase Auth 完成认证，加载用户角色和权限集合。", "用户已拥有账号且账号未被禁用。", "登录成功进入招聘指挥台；登录失败展示错误原因；权限加载失败时禁止进入业务页面。"),
        ("UC-02 创建招聘岗位", "招聘负责人录入岗位名称、部门、地点、技术要求、学历经验要求、筛选阈值和岗位状态。", "用户具备 MANAGE_POSITIONS 权限。", "系统保存岗位并在岗位列表、筛选下拉和候选人关联入口中可见。"),
        ("UC-03 上传并筛选简历", "招聘专员选择岗位后上传单份或多份简历，系统完成文件保存、文本提取、结构化解析和岗位匹配。", "岗位有效，文件类型和大小符合限制。", "系统生成上传记录、候选人记录、匹配分、推荐结论、证据和失败明细。"),
        ("UC-04 复核筛选结果", "招聘人员查看候选人与岗位的匹配证据，结合技能命中、缺失项、项目经历和风险点做人工结论。", "候选人和匹配结果已生成。", "人工复核结论保存到匹配记录，后续列表和详情使用最新人工状态展示。"),
        ("UC-05 创建面试安排", "招聘人员选择候选人、岗位、面试阶段、时间、面试官和房间类型创建面试。", "候选人已入库并关联岗位。", "系统生成 upcoming_interviews 记录，并可进入面试准备流程。"),
        ("UC-06 进入 AI 面试房间", "候选人通过面试房间入口进入在线面试，系统校验房间口令和面试状态。", "面试处于可进入状态，房间口令有效。", "系统创建或恢复 interview_sessions，并展示面试说明和开始入口。"),
        ("UC-07 Agent 多轮问答", "候选人启动面试后，Agent 基于简历和岗位生成问题，候选人逐轮回答，Agent 决定追问或下一题。", "Agent 服务可用，FastAPI 能组装完整上下文。", "每一轮候选人回答和 Agent 消息写入 interview_turns，页面刷新后可恢复。"),
        ("UC-08 生成面试报告", "面试结束后，系统从 Agent 获取 final_report，映射为面试报告并展示总分、维度分、优势、风险和建议。", "面试会话已结束或进入评分阶段。", "interview_reports 写入成功；若 Agent 未返回报告，生成待人工复核报告。"),
        ("UC-09 人工确认报告", "招聘人员阅读 AI 报告后确认、修正或补充最终建议。", "报告已生成且用户具备 MANAGE_INTERVIEWS 权限。", "系统保存人工确认人、确认时间、最终建议和备注。"),
        ("UC-10 导入薪酬市场数据", "管理员导入 CSV、JSON 或 JSONL 薪酬数据，系统保存原始记录并执行标准化和基准刷新。", "用户具备薪酬管理权限且服务密钥配置正确。", "系统生成原始记录、标准化记录、基准数据和导入任务统计。"),
        ("UC-11 管理用户权限", "超级管理员查看用户列表，设置角色和权限键。", "用户具备 owner 或 super_admin 角色。", "权限变更写入 user_roles，目标用户重新进入页面后权限生效。"),
        ("UC-12 调整公司 AI 设置", "管理员调整默认模型、AI 策略、简历隐私、面试题量和反馈要求。", "用户具备 MANAGE_SETTINGS 权限。", "设置写入 company_settings，后续简历筛选和面试流程读取最新配置。"),
    ]
    for uc_id, name, pre, result in use_cases:
        doc.add_heading(f"{uc_id} {name}", level=2)
        add_kv_table(
            doc,
            [
                ("用例描述", name),
                ("前置条件", pre),
                ("主成功场景", result),
                ("异常扩展", "登录失效、权限不足、关联数据缺失、外部服务失败或数据库写入失败时，系统保留当前上下文并给出可理解提示。"),
                ("验收标准", "用例在真实页面中可完整执行，数据库记录与页面结果一致，异常路径不产生脏数据或不可恢复状态。"),
            ],
        )

    doc.add_heading("第十一章 业务规则详细清单", level=1)
    rule_groups = [
        ("岗位规则", ["岗位名称不能为空。", "停用岗位不得创建新的筛选任务。", "筛选阈值影响推荐分组和人工复核优先级。", "岗位 JD 结构化结果需要与岗位主表关联保存。"]),
        ("简历规则", ["同一文件重复上传时应提示或复用已有处理结果。", "简历解析失败不得生成完整候选人主档。", "结构化简历允许字段缺失，但必须保留错误和置信度。", "匹配结果必须能回溯到岗位要求和简历证据。"]),
        ("候选人规则", ["候选人可关联多个岗位匹配结果。", "候选人删除需处理关联上传、匹配、面试和报告的业务影响。", "候选人联系方式属于敏感信息，应按权限展示。", "人工复核结论优先于 AI 推荐结论用于招聘推进。"]),
        ("面试规则", ["面试开始前必须校验候选人、岗位和会话。", "已取消或已结束面试不得重新进入进行中状态，除非管理员执行受控恢复。", "面试轮次按 session_id 和 turn_no 排序。", "最终面试结论必须支持人工确认。"]),
        ("Agent 规则", ["Agent 只能通过 FastAPI 间接接入。", "session_id 是 Agent 会话线程键。", "Agent 题量不得超过配置上限。", "Agent 报告缺失时进入 pending_human_review。"]),
        ("薪酬规则", ["原始薪酬数据必须先入库再标准化。", "标准化失败不得丢弃原始记录。", "样本量不足的基准需要标记置信风险。", "薪酬数据仅供决策参考，不作为系统自动录用条件。"]),
        ("权限规则", ["页面入口、接口调用和数据库策略必须一致。", "服务密钥不得进入浏览器端。", "首次超级管理员认领只能在系统无超级管理员时发生。", "权限变更需要记录操作者和目标用户。"]),
    ]
    for group, rules in rule_groups:
        doc.add_heading(group, level=2)
        add_bullets(doc, rules)

    doc.add_heading("第十二章 验收矩阵", level=1)
    matrix_rows = [
        ("登录权限", "普通用户无法进入管理员配置；管理员可管理岗位；超级管理员可维护权限。", "页面入口、接口返回、数据库访问均符合权限。"),
        ("岗位流程", "创建岗位、编辑岗位、停用岗位、筛选选择岗位。", "岗位状态在所有关联页面一致。"),
        ("简历流程", "上传、解析、匹配、失败重试、取消任务。", "每个阶段都有状态，失败有错误码和错误信息。"),
        ("候选人流程", "查看详情、人工复核、删除候选人、查看面试历史。", "详情完整且无越权字段暴露。"),
        ("AI 面试", "prepare、start、turn、score、human-confirm。", "会话、轮次和报告全部落库，可刷新恢复。"),
        ("Agent 异常", "Agent 超时、密钥错误、返回非 JSON、无 final_report。", "系统不丢失已回答内容，并进入可恢复或待复核状态。"),
        ("薪酬流程", "导入、标准化、基准刷新、决策看板。", "原始数据和标准化数据均可追踪。"),
        ("设置流程", "修改公司设置、模型配置和面试题量。", "后续流程读取新设置，历史记录不被错误覆盖。"),
    ]
    for scope, case, standard in matrix_rows:
        doc.add_heading(scope, level=2)
        doc.add_paragraph(f"验收场景：{case}")
        doc.add_paragraph(f"通过标准：{standard}")

    doc.add_heading("第十三章 项目交付清单", level=1)
    deliverables = [
        ("源代码交付", "交付内容包括前端 React/TypeScript 源码、FastAPI 后端源码、数据库迁移、导入脚本、配置样例和运行文档。源码需保持目录清晰，关键配置通过环境变量提供，不将生产密钥写入仓库。"),
        ("数据库交付", "交付内容包括 Supabase 迁移文件、RLS 策略、RPC 函数、表结构说明和必要初始化数据。交付时需说明迁移执行顺序、依赖关系和回滚注意事项。"),
        ("文档交付", "交付内容包括需求规格说明书、概要设计说明书、运行部署说明、权限说明、Agent 集成说明、AI 面试运行说明和薪酬市场操作说明。文档需与当前代码和数据库结构保持一致。"),
        ("测试交付", "交付内容包括功能验收记录、接口测试记录、端到端流程截图或日志、异常场景验证记录和未解决问题清单。测试结论应说明是否满足上线或演示要求。"),
        ("部署交付", "交付内容包括前端构建产物、后端启动方式、Agent 服务接入方式、环境变量清单、健康检查方式和发布后冒烟测试步骤。"),
        ("数据交付", "演示或验收环境需准备至少一个招聘岗位、多份简历、多个候选人、一次完整 AI 面试记录、一个人工确认报告和一组薪酬市场样例数据。"),
        ("运维交付", "交付内容包括常见故障处理方法、日志排查入口、Agent 不可用处理、数据库迁移失败处理、简历处理失败重试和权限配置恢复方式。"),
        ("最终确认", "项目验收完成后，应由业务负责人确认岗位、筛选、候选人、面试、Agent、薪酬、权限和设置模块均可按预期运行，并记录验收日期、验收人和遗留事项。"),
    ]
    for title, body in deliverables:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)


def add_design_extra_detail(doc: Document) -> None:
    doc.add_heading("第十七章 前端详细设计补充", level=1)
    frontend_sections = [
        ("路由组织", "前端使用 React Router 组织页面。受保护页面通过 ProtectedRoute 或等价机制校验登录态和权限键，未登录用户跳转登录页，权限不足用户展示受限提示。"),
        ("状态管理", "页面状态分为远程数据状态、表单输入状态、长任务状态和 UI 临时状态。长任务状态以数据库或后端返回为准，避免只保存在组件内存。"),
        ("错误展示", "接口错误需要转为用户可理解文案。上传、筛选、Agent、薪酬导入等长流程还需要展示错误码或错误摘要，便于用户反馈和后端排查。"),
        ("页面复用", "列表、详情、筛选条件、状态徽标、确认弹窗、空状态和错误提示应形成可复用组件，保证不同模块交互一致。"),
        ("截图与文档一致", "文档中的页面图例来自 ui-screenshots，后续 UI 大改时应重新截图并更新说明书。"),
    ]
    for title, body in frontend_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    doc.add_heading("第十八章 后端详细设计补充", level=1)
    backend_sections = [
        ("接口分层", "FastAPI 入口负责 HTTP 参数解析和响应；业务函数负责流程编排；数据访问函数负责 Supabase 查询和写入；外部服务函数负责 Agent、LLM、OCR 等调用。"),
        ("鉴权流程", "后端从请求中解析用户身份，调用 require_user 或等价逻辑校验登录态，再根据接口场景校验权限键和业务对象归属。"),
        ("Agent 调用封装", "agent_fetch 统一处理 AGENT_BASE_URL、x-agent-secret、超时、状态码和 JSON 解析。所有 Agent 调用都应通过该封装，避免重复错误处理。"),
        ("数据映射", "后端负责将 Agent、LLM 或导入脚本返回的数据映射为系统稳定字段。外部字段变化不应直接影响前端展示模型。"),
        ("事务边界", "Supabase REST/RPC 写入需要按业务原子性设计。无法单事务完成的长流程，应通过状态字段和幂等键保证重试安全。"),
    ]
    for title, body in backend_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    doc.add_heading("第十九章 Agent 服务约定补充", level=1)
    agent_contract = [
        ("启动请求", "POST /agent/start 接收 session_id、resume_text、jd_text、question_count，返回 interview_plan、message、status 和 state_snapshot。"),
        ("答题请求", "POST /agent/answer 接收 session_id 和 user_answer，返回下一条 message、status、state_snapshot，可附带 question_id 和 answer_guidance。"),
        ("状态请求", "GET /agent/status?session_id=... 返回当前状态、state_snapshot 和可选 final_report。"),
        ("报告结构", "final_report 至少包含 overall_score、hire_recommendation、strengths、concerns 或 risks；建议包含 dimension_scores 和 detailed_evaluations。"),
        ("错误结构", "Agent 错误建议返回 detail、error_code 和可恢复标识，FastAPI 将其转换为面试流程错误。"),
    ]
    for title, body in agent_contract:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
        add_bullets(
            doc,
            [
                "字段命名保持稳定，新增字段向后兼容。",
                "不得返回与 session_id 不匹配的其他候选人内容。",
                "敏感上下文只用于面试推理，不在 Agent 日志中完整明文保存。",
            ],
        )

    doc.add_heading("第二十章 数据安全与隐私补充", level=1)
    privacy_sections = [
        ("候选人隐私", "候选人简历、联系方式、面试回答和薪酬画像均属于敏感数据。页面展示应基于权限控制，日志中避免输出完整明文。"),
        ("模型调用隐私", "调用 LLM、OCR 或 Agent 时只传递完成任务所需上下文。记录用量时保留业务 ID、耗时和 token，不保存不必要的完整简历文本。"),
        ("权限数据安全", "user_roles、company_settings 和 llm_model_configs 的写入需限制管理员权限，普通用户不得直接修改。"),
        ("对象存储安全", "简历文件存储路径应避免泄露个人信息。下载或预览应通过授权 URL 或受控接口完成。"),
        ("人工复核责任", "AI 输出不作为最终录用决定。系统通过人工确认字段记录最终责任人和确认时间。"),
    ]
    for title, body in privacy_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)


def create_requirements_doc() -> Path:
    doc = new_doc_from_template("需求")
    add_template_cover(doc, "需求规格说明书", "面向智能招聘、简历筛选、AI 面试和权限运营的系统需求说明")
    add_revision_table(doc)
    add_toc(
        doc,
        [
            "第一章 引言",
            "1.1 编写目的",
            "1.2 文档范围",
            "1.3 项目概要",
            "1.4 术语和缩写",
            "1.5 参考资料",
            "第二章 任务概述",
            "2.1 目标",
            "2.2 用户的特点",
            "2.3 假定和约束",
            "第三章 需求定义",
            "3.1 功能需求",
            "3.1.10 AI 面试 Agent 专项需求",
            "3.2 性能需求",
            "3.3 输入输出需求",
            "3.4 数据管理能力需求",
            "3.5 故障处理需求",
            "第四章 运行环境",
            "第五章 详细功能规格",
            "第六章 数据字段与数据质量需求",
            "第七章 接口与集成需求扩展",
            "第八章 测试与验收需求",
            "第九章 运维与部署需求",
            "第十章 用户用例详细说明",
            "第十一章 业务规则详细清单",
            "第十二章 验收矩阵",
            "第十三章 项目交付清单",
        ],
    )
    doc.add_heading("第一章 引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    doc.add_paragraph(
        f"本文档用于明确{PROJECT_NAME}的业务目标、功能范围、非功能要求、输入输出、数据管理和运行环境，为后续概要设计、详细设计、开发、测试、部署和验收提供依据。"
    )
    doc.add_heading("1.2 文档范围", level=2)
    doc.add_paragraph("本文档覆盖 Web 前端、FastAPI 业务后端、Supabase 数据与存储、AI 简历筛选、AI 面试、薪酬市场数据处理、权限与系统设置等功能。")
    doc.add_heading("1.3 项目概要", level=2)
    add_bullets(
        doc,
        [
            "系统名称：RecruitPro智能招聘管理系统。",
            "系统定位：面向招聘团队的人才筛选、面试协同和招聘决策平台。",
            "系统形态：B/S 架构 Web 应用，前端采用 React + TypeScript + Vite，业务后端采用 FastAPI，数据层采用 Supabase/PostgreSQL。",
            "核心价值：将岗位要求、简历解析、候选人匹配、AI 面试、人工复核、薪酬参考和权限治理整合到统一工作台。",
        ],
    )
    doc.add_heading("1.4 术语和缩写", level=2)
    terms = [
        ("1", "RLS", "Row Level Security，数据库行级安全策略，用于限制用户只能访问被授权的数据。"),
        ("2", "LLM", "大语言模型，用于简历结构化、岗位匹配、面试问答和评分等 AI 能力。"),
        ("3", "OCR", "光学字符识别，用于在简历文本无法直接提取时进行备用解析。"),
        ("4", "Agent", "外部智能体服务，负责 AI 面试过程中的问题生成、追问和报告生成。"),
        ("5", "匹配分", "候选人与岗位要求之间的综合评分，包含技能、项目、经验、学历和证据维度。"),
    ]
    for no, name, desc in terms:
        doc.add_paragraph(f"{no}. {name}：{desc}")
    doc.add_heading("1.5 参考资料", level=2)
    add_bullets(
        doc,
        [
            "项目源码：D:/project/RecruitPro_。",
            "数据库迁移：supabase/migrations。",
            "运行文档：docs/phase1-pipeline.md、docs/ai-interview-runtime.md、docs/agent-integration-runtime.md、docs/permissions.md、docs/salary-market-ops.md。",
            "UI 截图：ui-screenshots 目录。",
        ],
    )

    doc.add_heading("第二章 任务概述", level=1)
    doc.add_heading("2.1 目标", level=2)
    add_bullets(
        doc,
        [
            "建立从岗位发布、简历上传、AI 解析、候选人匹配、人工复核到面试推进的闭环流程。",
            "通过可解释分数、匹配证据、风险提示和人工决策字段提升筛选过程透明度。",
            "支持 AI 面试房间、面试轮次记录、自动评分和人工确认，减少重复沟通成本。",
            "以权限键、角色层级和数据库 RLS 控制业务边界，保证招聘数据安全。",
            "引入薪酬市场数据导入与基准刷新能力，辅助岗位预算和 offer 决策。",
        ],
    )
    doc.add_heading("2.2 用户的特点", level=2)
    add_bullets(
        doc,
        [
            "招聘负责人：关注岗位进度、候选人质量、面试安排和最终录用建议。",
            "HR/招聘专员：负责简历上传、筛选结果复核、候选人推进和面试邀约。",
            "面试官：关注候选人项目证据、面试问题、评分报告和风险点。",
            "系统管理员：负责组织设置、AI 策略、账号权限、模型配置和数据治理。",
            "候选人：通过面试房间参与在线 AI 面试。候选人侧仅接触面试相关入口。",
        ],
    )
    doc.add_heading("2.3 假定和约束", level=2)
    add_bullets(
        doc,
        [
            "系统需要可访问 Supabase 项目、Storage、PostgreSQL 数据库和必要 RPC。",
            "AI 简历筛选和 AI 面试依赖可用的模型服务、OCR 能力和外部 Agent 服务。",
            "生产环境需配置 VITE_SUPABASE_URL、VITE_SUPABASE_ANON_KEY、SUPABASE_SERVICE_ROLE_KEY、AGENT_BASE_URL、AGENT_SHARED_SECRET 等环境变量。",
            "系统以浏览器为主要访问方式，建议部署于 HTTPS 环境并限制服务密钥只在后端使用。",
            "用户角色和权限变更必须通过受控 RPC 或后端接口完成，不允许普通客户端直接修改敏感表。",
        ],
    )

    doc.add_heading("第三章 需求定义", level=1)
    doc.add_heading("3.1 功能需求", level=2)
    add_module_requirements(doc)
    add_agent_requirements(doc)
    doc.add_heading("3.2 性能需求", level=2)
    add_bullets(
        doc,
        [
            "常规列表查询和详情查询在正常网络与数据库负载下应在 3 秒内返回。",
            "前端交互响应应在 1 秒内给出状态变化或加载反馈。",
            "批量简历处理支持任务状态追踪、失败记录、取消和重试，避免长任务无反馈。",
            "AI 和 OCR 等外部服务调用需设置超时，失败时写入明确错误码和错误信息。",
        ],
    )
    doc.add_heading("3.3 输入输出需求", level=2)
    rows = [
        ("账号登录", "邮箱、密码、重置密码链接", "登录态、权限集合、错误提示"),
        ("岗位管理", "岗位基础信息、筛选阈值、技术要求", "岗位列表、岗位详情、筛选规则"),
        ("简历筛选", "PDF/DOCX 文件、目标岗位", "结构化简历、匹配分、证据、候选人记录"),
        ("面试管理", "候选人、时间、面试官、房间口令", "面试安排、房间链接、面试报告"),
        ("薪酬导入", "CSV/JSON/JSONL 市场薪酬数据", "标准化薪酬记录、基准数据、导入任务状态"),
    ]
    for category, input_text, output_text in rows:
        doc.add_paragraph(f"{category}：输入为{input_text}；输出为{output_text}。")
    doc.add_heading("3.4 数据管理能力需求", level=2)
    add_bullets(
        doc,
        [
            "系统需管理岗位、候选人、简历上传、结构化简历、岗位要求、匹配结果、面试安排、面试会话、问答轮次、面试报告、薪酬市场数据、模型配置、公司设置和用户权限。",
            "简历原文件保存至 Supabase Storage，结构化结果和处理状态保存至 PostgreSQL。",
            "关键业务表启用 RLS；后台服务使用服务密钥进行受控写入，前端按权限读取和操作。",
            "AI 输出应保留原始载荷、置信度、证据片段和人工复核字段，以支持追溯。",
        ],
    )
    doc.add_heading("3.5 故障处理需求", level=2)
    add_bullets(
        doc,
        [
            "上传、OCR、LLM、数据库写入、权限校验、外部 Agent 调用失败时需返回可读错误信息。",
            "简历处理流程需记录 pipeline_stage、status、error_code、error_message 和 retry_count。",
            "面试流程异常中断时应保留已产生的会话、轮次和报告状态，便于人工处理。",
            "数据库迁移缺失或结构不一致时，系统应提示同步迁移，不应静默失败。",
        ],
    )

    doc.add_heading("第四章 运行环境", level=1)
    doc.add_heading("4.1 设备", level=2)
    doc.add_paragraph("客户端为现代浏览器；服务端可部署于 Linux 或容器环境；数据库和对象存储依赖 Supabase。")
    doc.add_heading("4.2 支持软件", level=2)
    add_bullets(
        doc,
        [
            "前端：React 19、TypeScript、Vite、Tailwind CSS、React Router、Lucide React、Supabase JS。",
            "后端：FastAPI、Python、Supabase Python Client、httpx。",
            "数据库：PostgreSQL/Supabase，包含 SQL 迁移、RLS 策略、RPC 和 Storage。",
            "AI 服务：外部 LLM/OCR 服务与 Agent 服务，面试 Agent 通过 AGENT_BASE_URL 接入。",
        ],
    )
    doc.add_heading("4.3 接口", level=2)
    doc.add_paragraph("系统接口主要包括前端到 Supabase 的数据访问接口、前端到 FastAPI 的业务接口、FastAPI 到 Supabase 的服务端接口、FastAPI 到外部 Agent/LLM/OCR 的集成接口。")
    doc.add_heading("4.4 控制", level=2)
    doc.add_paragraph("系统通过浏览器人机交互控制业务流程，通过权限键控制页面和操作可见性，通过后台接口与数据库策略控制敏感数据写入。")
    doc.add_heading("第四章补充说明", level=2)
    add_bullets(
        doc,
        [
            "本地开发环境建议前端运行 Vite 服务，后端运行 FastAPI 服务，数据库连接 Supabase 项目或本地 Supabase 栈。",
            "生产部署环境建议前端静态资源、FastAPI 后端、Agent 服务、Supabase 数据库分别独立部署并通过环境变量连接。",
            "Agent 服务应与 FastAPI 保持网络可达，生产环境建议只允许后端服务访问 Agent。",
            "数据库迁移文件是结构来源，部署前必须确认迁移已应用，否则简历、面试和权限模块可能出现字段缺失。",
            "模型服务和 OCR 服务属于外部依赖，应在系统设置或后端环境中配置，失败时需要可观测日志。",
        ],
    )
    add_requirements_deep_expansion(doc)
    add_requirements_full_appendix(doc)
    add_requirements_extra_detail(doc)

    path = OUT_DIR / "RecruitPro智能招聘管理系统需求规格说明书.docx"
    normalize_body_alignment(doc)
    doc.save(path)
    return path


def create_design_doc() -> Path:
    doc = new_doc_from_template("概要")
    add_template_cover(doc, "概要设计说明书", "系统架构、模块设计、接口设计、数据结构与部署说明")
    add_revision_table(doc)
    add_toc(
        doc,
        [
            "第一章 引言",
            "第二章 总体设计",
            "2.1 需求规定",
            "2.2 运行环境",
            "2.3 基本设计概念和处理流程",
            "2.4 系统架构",
            "2.5 功能模块设计",
            "第三章 接口设计",
            "3.4 Agent 集成接口",
            "第四章 运行设计",
            "4.4 Agent 运行流程",
            "第五章 系统数据结构设计",
            "第六章 系统安全设计",
            "第七章 部署与运维设计",
            "第八章 模块详细设计",
            "第九章 Agent 详细设计",
            "第十章 数据库详细设计",
            "第十一章 错误处理与日志设计",
            "第十二章 测试设计",
            "第十三章 接口详细设计补充",
            "第十四章 部署架构详细设计",
            "第十五章 数据迁移与版本管理设计",
            "第十六章 可维护性与扩展设计",
            "第十七章 前端详细设计补充",
            "第十八章 后端详细设计补充",
            "第十九章 Agent 服务约定补充",
            "第二十章 数据安全与隐私补充",
        ],
    )
    doc.add_heading("第一章 引言", level=1)
    doc.add_paragraph(f"本文档在需求规格说明基础上，对{PROJECT_NAME}的总体架构、模块职责、接口、数据结构、安全和部署方式进行概要设计说明。")
    doc.add_heading("项目概要", level=2)
    doc.add_paragraph("RecruitPro 是面向招聘业务的 Web 应用，围绕岗位、简历、候选人、面试、薪酬和权限形成闭环。系统采用前后端分离，前端直接使用 Supabase 读取部分授权数据，复杂业务和服务密钥操作集中在 FastAPI 后端。")
    doc.add_heading("术语和缩写", level=2)
    add_kv_table(
        doc,
        [
            ("Frontend", "React + TypeScript + Vite 前端应用。"),
            ("Backend", "FastAPI 业务后端，负责简历处理、面试流程、管理员操作和外部服务集成。"),
            ("Supabase", "提供 PostgreSQL、Auth、Storage、RLS 和 RPC 的数据平台。"),
            ("Agent", "外部 AI 面试服务，使用共享密钥与 FastAPI 通信。"),
            ("Pipeline", "简历上传后的阶段化处理链路。"),
        ],
    )

    doc.add_heading("第二章 总体设计", level=1)
    doc.add_heading("2.1 需求规定", level=2)
    doc.add_paragraph("系统需满足招聘业务在岗位维护、简历筛选、候选人推进、AI 面试、薪酬参考、权限管理和系统配置方面的主要业务需求，并保证关键流程可追踪、可复核、可恢复。")
    doc.add_heading("2.2 运行环境", level=2)
    add_kv_table(
        doc,
        [
            ("浏览器端", "现代桌面浏览器，访问 React 单页应用。"),
            ("前端构建", "Vite + TypeScript，构建产物部署为静态资源。"),
            ("业务后端", "FastAPI，默认本地开发端口为 8010，可容器化部署。"),
            ("数据服务", "Supabase PostgreSQL、Auth、Storage、RPC、RLS。"),
            ("外部服务", "LLM/OCR 服务、AI 面试 Agent 服务。"),
        ],
    )
    doc.add_heading("2.3 基本设计概念和处理流程", level=2)
    doc.add_heading("系统设计概念", level=3)
    add_bullets(
        doc,
        [
            "以岗位为业务中心：岗位规则驱动简历筛选、候选人匹配和面试推进。",
            "以证据为 AI 输出基础：AI 解析和评分需要保留证据片段、结构化字段、置信度和人工复核状态。",
            "以前后端分离为工程边界：展示和交互在前端，长流程和敏感操作在后端。",
            "以数据库安全为底线：权限不只依赖前端路由，还需通过 RLS、RPC 和后端接口共同约束。",
        ],
    )
    doc.add_heading("系统处理流程", level=3)
    flows = [
        ("简历筛选流程", "screening 页面上传文件", "创建上传记录、文件入库、文本提取、结构化解析、岗位匹配、候选人持久化", "候选人、匹配分、证据和复核状态"),
        ("AI 面试流程", "interviews 页面创建或进入面试", "准备面试、创建会话、轮次问答、结束面试、评分、人工确认", "面试报告、维度分、建议和风险"),
        ("权限管理流程", "settings/access 页面", "超级管理员通过受控接口修改角色和权限", "user_roles 更新并影响页面访问"),
        ("薪酬导入流程", "导入脚本或接口", "读取文件、字段映射、原始入库、标准化、基准刷新、任务记录", "薪酬市场基准和候选人薪酬参考"),
    ]
    for name, trigger, process, output in flows:
        doc.add_heading(name, level=4)
        add_kv_table(
            doc,
            [
                ("触发入口", trigger),
                ("主要处理", process),
                ("输出结果", output),
            ],
        )

    doc.add_heading("2.4 系统架构", level=2)
    rows = [
        ("表现层", "React SPA、React Router、Tailwind CSS", "页面路由、权限保护、数据展示、文件上传、用户交互。"),
        ("业务层", "FastAPI backend/main.py", "简历筛选、薪酬导入、面试运行、管理员接口、外部 Agent 调用。"),
        ("数据层", "Supabase PostgreSQL", "岗位、候选人、简历、匹配、面试、薪酬、权限、配置等结构化数据。"),
        ("存储层", "Supabase Storage", "简历文件、头像等对象存储。"),
        ("AI 服务层", "LLM/OCR/Agent", "简历解析、岗位匹配、面试问答和评分报告生成。"),
    ]
    for layer, component, duty in rows:
        doc.add_paragraph(f"{layer}：{component}。{duty}")
    doc.add_heading("2.5 功能模块设计", level=2)
    module_rows = [
        ("认证与路由", "Login、ResetPassword、ProtectedRoute", "Supabase Auth、user_roles", "登录态和权限键共同决定可访问页面。"),
        ("招聘指挥台", "Dashboard", "active_positions、candidates、upcoming_interviews", "聚合展示关键业务数据并提供快速入口。"),
        ("岗位管理", "Positions", "active_positions、parsed_job_requirements", "岗位规则作为筛选和匹配基础。"),
        ("简历筛选", "Screening、screeningPipeline", "resume_uploads、parsed_*、candidate_position_matches、candidates", "阶段化处理，支持批量、失败、取消和人工复核。"),
        ("候选人详情", "Candidates、CandidateDetail", "candidates、matches、projects、profiles", "展示匹配解释、证据、风险和人工结论。"),
        ("面试管理", "Interviews、InterviewRoom、interviewRuntime", "upcoming_interviews、interview_sessions、turns、reports", "FastAPI 对接 Agent，记录全过程并生成报告。"),
        ("系统设置", "Settings、OrganizationSettings、AiPolicySettings、AdminManagement", "company_settings、llm_model_configs、user_roles", "组织、模型、面试策略和权限集中管理。"),
        ("薪酬决策", "Salary 相关视图与导入脚本", "market_salary_*、candidate_salary_profiles", "支持市场数据清洗、基准刷新和候选人薪酬画像。"),
    ]
    for name, frontend, backend, point in module_rows:
        doc.add_heading(name, level=3)
        add_kv_table(
            doc,
            [
                ("前端页面/组件", frontend),
                ("后端/数据依赖", backend),
                ("设计要点", point),
            ],
        )

    doc.add_heading("2.6 模块内部设计补充", level=2)
    for name, frontend, backend, point in module_rows:
        doc.add_heading(name, level=3)
        add_bullets(
            doc,
            [
                f"前端组成：{frontend}。",
                f"后端和数据依赖：{backend}。",
                f"关键设计：{point}",
                "状态管理：页面需区分加载中、空数据、处理成功、处理失败和权限不足状态。",
                "数据一致性：用户提交后应以数据库返回结果刷新页面，不应只依赖本地乐观状态。",
                "扩展点：后续可在保持数据结构兼容的前提下扩展字段、筛选条件、Agent 策略或报表维度。",
            ],
        )

    doc.add_heading("2.7 Agent 组件设计", level=2)
    add_bullets(
        doc,
        [
            "Agent 网关组件位于 FastAPI 后端，统一封装 agent_fetch，避免前端直接访问 Agent 服务。",
            "上下文构建组件负责将候选人、结构化简历、项目经历、岗位和结构化 JD 拼接为 Agent 可理解的文本。",
            "问题计划映射组件负责将 Agent interview_plan 转换为前端可展示和数据库可持久化的 question_plan。",
            "轮次记录组件负责将候选人回答和 AI 输出依次写入 interview_turns，并记录 source=agent 的元数据。",
            "报告映射组件负责将 final_report 转换为 interview_reports 的分数、建议、风险、证据和摘要字段。",
            "人工确认组件负责把最终人工判断覆盖或确认到报告中，保证 AI 建议不直接等同最终决策。",
        ],
    )

    doc.add_heading("第三章 接口设计", level=1)
    doc.add_heading("3.1 前端路由接口", level=2)
    add_bullets(
        doc,
        [
            "/login：登录页面。",
            "/reset-password：密码重置回调页面。",
            "/：招聘指挥台，需要 VIEW_DASHBOARD。",
            "/positions：岗位管理，需要 MANAGE_POSITIONS。",
            "/screening：简历筛选，需要 SCREEN_RESUMES。",
            "/candidates 与 /candidates/:id：候选人列表和详情，需要 VIEW_CANDIDATES。",
            "/interviews 与 /interview-room/:interviewId：面试管理和面试房间，需要 MANAGE_INTERVIEWS 或登录态。",
            "/settings/*：组织设置、AI 策略和权限管理，需要 MANAGE_SETTINGS，权限页需要超级管理员。",
        ],
    )
    doc.add_heading("3.2 后端 API 接口", level=2)
    api_groups = [
        ("健康检查", "GET /api/health"),
        ("配置与薪酬", "GET/PATCH /api/settings/company，GET /api/salaries，GET /api/salary/dashboard，POST /api/salary/market/import，POST /api/salary/market/refresh"),
        ("管理员", "GET /api/admin/has-super-admin，POST /api/admin/claim-initial-super-admin，GET /api/admin/users，POST /api/admin/users/{id}/permissions，POST /api/admin/users/{id}/role"),
        ("岗位与筛选", "GET/POST/PATCH/DELETE /api/positions，POST /api/screening/phase1，POST /api/screening/persist-phase1，POST /api/screening/rescreen"),
        ("候选人与上传", "GET /api/candidates，GET /api/candidates/{id}/detail，DELETE /api/candidates/{id}，POST /api/uploads，PATCH /api/uploads/{id}，POST /api/uploads/delete"),
        ("面试", "GET/POST/PATCH/DELETE /api/interviews，POST /api/interviews/prepare，POST /api/interviews/start，POST /api/interviews/turn，POST /api/interviews/finish，POST /api/interviews/score，POST /api/interviews/human-confirm"),
    ]
    add_kv_table(doc, api_groups)
    doc.add_heading("3.3 外部接口", level=2)
    add_bullets(
        doc,
        [
            "Supabase Auth：登录、密码重置、用户身份识别。",
            "Supabase Storage：简历文件和头像等对象存储。",
            "外部 Agent：通过 AGENT_BASE_URL 和 AGENT_SHARED_SECRET 接入 AI 面试服务。",
            "LLM/OCR：通过后端配置的模型服务完成文本识别、结构化和评分。",
        ],
    )
    doc.add_heading("3.4 Agent 集成接口", level=2)
    doc.add_paragraph(
        "RecruitPro 当前不再使用 Supabase Edge Functions 编排 AI 面试，而是由前端调用 FastAPI，FastAPI 再调用外部 Agent 服务。Agent 网关统一由 backend/main.py 中的 agent_fetch 处理，负责拼接 AGENT_BASE_URL、附加 x-agent-secret、设置超时并解析 JSON 返回。"
    )
    add_kv_table(
        doc,
        [
            ("启动接口", "FastAPI /api/interviews/start 调用 Agent /agent/start，传入 session_id、resume_text、jd_text、question_count，接收 interview_plan 与首条 message。"),
            ("答题接口", "FastAPI /api/interviews/turn 调用 Agent /agent/answer，传入 session_id 和 user_answer，接收下一条 message、status 与 state_snapshot。"),
            ("状态接口", "FastAPI /api/interviews/score 调用 Agent /agent/status?session_id=...，读取 final_report 并映射为面试报告。"),
            ("鉴权方式", "FastAPI 在请求头中写入 x-agent-secret，Agent 服务端应校验共享密钥。"),
            ("超时策略", "AGENT_TIMEOUT_MS 控制 Agent 请求超时时间，默认按毫秒配置并在后端转换为秒。"),
        ],
    )
    doc.add_heading("Agent 请求数据结构", level=3)
    add_bullets(
        doc,
        [
            "session_id：面试会话 ID，同时作为 Agent 线程键，用于串联多轮问答。",
            "resume_text：由候选人信息、结构化简历和项目经历拼接形成的候选人上下文。",
            "jd_text：由岗位信息和结构化岗位要求拼接形成的职位上下文。",
            "question_count：来自公司设置或面试配置的题量约束。",
            "user_answer：候选人在面试房间提交的单轮回答文本。",
        ],
    )
    doc.add_heading("3.5 后端接口详细说明", level=2)
    endpoint_details = [
        ("GET /api/dashboard", "聚合招聘指挥台数据，供首页展示岗位、候选人和面试概览。"),
        ("GET /api/screening/dashboard", "返回筛选模块仪表盘数据，支撑强匹配、待评估和淘汰分组。"),
        ("POST /api/screening/job-requirement/resolve", "根据岗位信息生成或解析结构化岗位要求。"),
        ("PATCH /api/positions/{position_id}/screening-review", "确认岗位筛选规则已复核，清除重筛提醒。"),
        ("PATCH /api/matches/{match_id}/manual-review", "写入候选人与岗位匹配结果的人工复核结论。"),
        ("POST /api/candidates/delete-batch", "批量删除候选人，需保证关联数据处理策略明确。"),
        ("POST /api/uploads/mark-failed", "将卡住或异常的上传任务批量标记失败。"),
        ("POST /api/uploads/{upload_id}/fail", "将单个上传任务置为失败并记录错误码。"),
        ("POST /api/uploads/{upload_id}/cancel", "取消尚未完成的上传或处理任务。"),
        ("GET /api/interviews/{interview_id}/report", "获取指定面试的最新 AI 报告。"),
        ("GET /api/interviews/sessions/{session_id}/turns", "按 turn_no 返回面试问答轮次。"),
        ("GET /api/interviews/candidates/{candidate_id}/position", "查询候选人关联岗位，为面试准备上下文。"),
        ("POST /api/llm-usage", "记录模型调用统计，用于成本、延迟和失败率分析。"),
    ]
    for endpoint, desc in endpoint_details:
        doc.add_paragraph(f"{endpoint}：{desc}")
    doc.add_heading("Agent 返回数据结构", level=3)
    add_bullets(
        doc,
        [
            "interview_plan：问题计划，系统将其映射为 interview_sessions.question_plan。",
            "message：Agent 给候选人的首问、追问、下一题或结束提示。",
            "status：Agent 当前状态，例如继续、等待复核或结束。",
            "state_snapshot：Agent 的运行快照，包含已问问题数、回答数、后续节点等。",
            "final_report：最终评价报告，系统映射为 interview_reports。",
        ],
    )

    doc.add_heading("第四章 运行设计", level=1)
    doc.add_heading("4.1 运行控制", level=2)
    add_bullets(
        doc,
        [
            "前端启动后根据登录态和权限加载页面；未登录用户重定向到登录页。",
            "简历筛选、AI 面试和薪酬导入等长流程通过后端状态字段反馈进度。",
            "面试房间使用面试 ID 和房间口令控制进入，并由后端维护会话状态。",
        ],
    )
    doc.add_heading("4.2 运行状态", level=2)
    add_kv_table(
        doc,
        [
            ("简历上传状态", "pending、processing、completed、failed 等。"),
            ("处理阶段", "uploaded、text_extraction、profile_extraction、matching、completed、failed。"),
            ("面试状态", "scheduled、ready、in_progress、completed、cancelled、no_show、failed。"),
            ("人工复核", "pass、pending、reject。"),
        ],
    )
    doc.add_heading("4.3 异常处理", level=2)
    doc.add_paragraph("系统将异常分为用户输入错误、权限错误、文件/存储错误、数据库结构错误、外部服务错误和网络超时错误。前端展示可理解提示，后端记录错误码和错误信息，关键流程保留可重试状态。")
    doc.add_heading("4.4 Agent 运行流程", level=2)
    add_bullets(
        doc,
        [
            "准备阶段：/api/interviews/prepare 校验用户、候选人、岗位和面试记录，创建或确认 interview_sessions，状态为 preparing/ready。",
            "启动阶段：/api/interviews/start 读取候选人、结构化简历、项目经历、岗位和结构化 JD，组装上下文后调用 /agent/start。",
            "首问落库：Agent 返回首条 message 后，系统以 speaker=ai 写入 interview_turns，并在 metadata 中记录 kind=question、question_id=agent-1、source=agent。",
            "答题阶段：候选人提交回答后，系统先写入 candidate turn，再调用 /agent/answer 获取下一条 AI message。",
            "追问/下一题：系统根据 state_snapshot.asked_question_count 判断 AI message 是下一题还是追问，并写入 kind=question/followup/closing。",
            "结束阶段：当 Agent status 为 wait_for_review 或 finish 时，系统将会话推进到 scoring 或 done，并提示进入评分。",
            "评分阶段：/api/interviews/score 查询 /agent/status，如果 final_report 已存在则映射并 upsert 到 interview_reports；如果不存在则生成 pending_human_review 报告。",
            "人工确认：/api/interviews/human-confirm 将人工最终结论写入报告，并回填 interview 的 ai_report_id 和状态。",
        ],
    )

    doc.add_heading("第五章 系统数据结构设计", level=1)
    doc.add_heading("5.1 主要数据表", level=2)
    data_rows = [
        ("权限", "user_roles", "用户角色、权限键和管理员层级。"),
        ("岗位", "active_positions、parsed_job_requirements", "岗位基础信息和结构化岗位要求。"),
        ("候选人", "candidates、candidate_position_matches、candidate_salary_profiles", "候选人档案、岗位匹配和薪酬画像。"),
        ("简历", "resume_uploads、parsed_resume_profiles、parsed_resume_projects", "上传状态、结构化简历和项目经历。"),
        ("面试", "upcoming_interviews、interview_sessions、interview_turns、interview_reports", "排期、会话、问答轮次和评分报告。"),
        ("薪酬", "market_salaries、market_salary_raw_records、market_salary_normalized_records、market_salary_benchmarks、market_salary_crawl_jobs", "薪酬市场数据、标准化数据、基准和导入任务。"),
        ("配置", "company_settings、llm_model_configs、llm_usage_events", "公司设置、模型配置和模型使用记录。"),
        ("会话", "chat_conversations、chat_messages", "招聘辅助会话数据。"),
    ]
    for domain, tables, desc in data_rows:
        doc.add_paragraph(f"{domain}：{tables}。{desc}")
    doc.add_heading("5.2 数据关系", level=2)
    add_bullets(
        doc,
        [
            "active_positions 与 candidates、resume_uploads、candidate_position_matches 存在岗位关联。",
            "resume_uploads 是简历处理链路入口，后续关联结构化简历、项目和匹配结果。",
            "candidate_position_matches 保存候选人与岗位的可解释匹配结果，候选人详情优先展示最新或指定匹配记录。",
            "upcoming_interviews 可关联 candidates，interview_sessions、interview_turns 和 interview_reports 共同构成面试闭环。",
            "market_salary_* 表共同支撑薪酬市场数据导入、标准化和基准刷新。",
        ],
    )
    doc.add_heading("5.3 Agent 相关数据结构", level=2)
    add_kv_table(
        doc,
        [
            ("interview_sessions.question_plan", "保存 Agent 生成的问题计划。每个问题可包含 question_id、topic、prompt、expected_signals、answer_guidance 等字段。"),
            ("interview_sessions.context_payload", "保存面试上下文摘要，可用于记录候选人、岗位、题量和运行模式等信息。"),
            ("interview_turns.metadata", "保存 Agent 轮次元数据，例如 kind、source、question_id、asked_question_count、answer_count、next_nodes、answer_guidance。"),
            ("interview_reports.dimension_scores", "保存 Agent 报告映射后的维度分，例如岗位匹配度、技术深度、项目证据、问题解决、沟通表达、主导力等。"),
            ("interview_reports.evidence", "保存 Agent 评分证据、状态快照或待复核说明，用于支持人工复核。"),
            ("interview_reports.human_confirmed", "标识 AI 报告是否经过人工确认，配套记录确认人和确认时间。"),
        ],
    )
    doc.add_heading("5.4 Agent 报告映射规则", level=2)
    add_bullets(
        doc,
        [
            "final_report.overall_score 映射为 interview_reports.overall_score。",
            "final_report.hire_recommendation 经过 map_agent_recommendation 归一化为 hire、hold、needs_review、reject。",
            "final_report.strengths 映射为 strengths，同时用于报告摘要。",
            "final_report.concerns 或 risks 映射为 risks，并参与 risk_score 计算。",
            "final_report.detailed_evaluations 映射为 question_evaluations，用于展示单题问题、回答、反馈、缺失逻辑点和分项评分。",
            "当 final_report 不存在时，系统写入 pending_human_review，证据中保留 agent_status 和 state_snapshot。",
        ],
    )
    doc.add_heading("5.5 数据表设计明细", level=2)
    table_designs = [
        ("active_positions", "岗位主表，保存岗位名称、部门、地点、状态、阈值、技术要求、年龄、学历和经验要求。"),
        ("candidates", "候选人主表，保存姓名、职位、经验、学历、年龄、匹配分、标签、高亮、城市和关联岗位。"),
        ("resume_uploads", "简历上传表，保存文件名、路径、大小、hash、状态、pipeline_stage、错误码、错误信息和重试次数。"),
        ("parsed_resume_profiles", "结构化简历画像表，保存基础画像、教育经历、技能、置信度和 LLM 原始输出。"),
        ("parsed_resume_projects", "候选人项目经历表，保存项目名称、摘要、技术栈、复杂度和职责。"),
        ("parsed_job_requirements", "结构化岗位要求表，保存岗位必须项、技能要求、经验学历要求和解析载荷。"),
        ("candidate_position_matches", "岗位匹配结果表，保存总分、推荐结论、技能匹配、缺失项、证据、需求分解和人工复核。"),
        ("upcoming_interviews", "面试安排表，保存候选人、阶段、岗位、时间、面试官、房间类型、状态和报告关联。"),
        ("interview_sessions", "面试会话表，保存会话模式、状态、question_plan、context_payload、开始和结束时间。"),
        ("interview_turns", "面试轮次表，保存轮次号、说话人、内容、输入模式、延迟、token、置信度和 metadata。"),
        ("interview_reports", "面试报告表，保存总分、维度分、优势、风险、建议、证据、摘要、风险分和人工确认信息。"),
        ("market_salary_raw_records", "薪酬原始记录表，保存来源职位、城市、薪酬文本、公司、采集时间和原始载荷。"),
        ("market_salary_normalized_records", "薪酬标准化记录表，保存清洗后的薪酬区间、币种、周期和标准职位城市。"),
        ("market_salary_benchmarks", "薪酬基准表，保存岗位城市维度的市场薪酬统计结果。"),
        ("user_roles", "用户角色权限表，保存角色、权限键和用户邮箱。"),
        ("company_settings", "公司设置表，保存 AI 策略、简历隐私、反馈要求、默认模型和面试配置。"),
        ("llm_model_configs", "模型配置表，保存模型供应商、模型名、协议和运行配置。"),
        ("llm_usage_events", "模型使用记录表，保存模型调用场景、token、延迟、成功状态和错误信息。"),
    ]
    for table_name, desc in table_designs:
        doc.add_paragraph(f"{table_name}：{desc}")

    doc.add_heading("第六章 系统安全设计", level=1)
    add_bullets(
        doc,
        [
            "认证采用 Supabase Auth，前端页面由 ProtectedRoute 控制访问。",
            "授权采用权限键：VIEW_DASHBOARD、MANAGE_POSITIONS、SCREEN_RESUMES、VIEW_CANDIDATES、MANAGE_INTERVIEWS、VIEW_SALARY、MANAGE_SETTINGS。",
            "角色层级采用 owner、super_admin、admin、user；高权限操作通过受控 RPC 或后端接口完成。",
            "关键业务表启用 RLS，普通客户端不能直接越权读写。",
            "服务密钥只允许后端和离线导入脚本使用，不应暴露在浏览器端。",
            "Agent 接口通过共享密钥和超时控制降低未授权调用和长时间阻塞风险。",
        ],
    )
    doc.add_heading("Agent 安全设计", level=2)
    add_bullets(
        doc,
        [
            "Agent 服务不直接暴露给前端页面，前端只能通过 FastAPI 业务接口进入面试流程。",
            "FastAPI 到 Agent 的请求使用 x-agent-secret 共享密钥，避免未授权服务伪造面试会话。",
            "Agent 不持有 Supabase service role key，不直接写入招聘数据库，降低数据泄露面。",
            "session_id 只作为 Agent 线程键，不应包含候选人隐私明文。",
            "FastAPI 负责对候选人、岗位、面试和用户权限做前置校验，Agent 只负责问答和评估逻辑。",
            "Agent 返回内容需经过后端映射和前端展示约束，最终录用建议必须支持人工确认。",
        ],
    )

    doc.add_heading("第七章 部署与运维设计", level=1)
    doc.add_heading("7.1 部署方式", level=2)
    add_bullets(
        doc,
        [
            "前端执行 npm run build 后部署 dist 静态资源。",
            "FastAPI 后端可直接运行或通过 backend/Dockerfile 容器化部署。",
            "docker-compose.yml 可作为本地或测试环境的组合部署入口。",
            "数据库结构以 supabase/migrations 为来源，通过 supabase db push 或 supabase db reset 同步。",
        ],
    )
    doc.add_heading("7.2 环境变量", level=2)
    add_kv_table(
        doc,
        [
            ("前端", "VITE_SUPABASE_URL、VITE_SUPABASE_ANON_KEY、VITE_API_BASE_URL。"),
            ("后端", "SUPABASE_URL、SUPABASE_SERVICE_ROLE_KEY、AGENT_BASE_URL、AGENT_SHARED_SECRET、AGENT_TIMEOUT_MS。"),
            ("导入脚本", "SUPABASE_URL 或 VITE_SUPABASE_URL、SUPABASE_SERVICE_ROLE_KEY。"),
        ],
    )
    doc.add_heading("7.3 Agent 部署要求", level=2)
    add_bullets(
        doc,
        [
            "Agent 服务需部署为可被 FastAPI 访问的 HTTP 服务，并提供 /agent/start、/agent/answer、/agent/status 接口。",
            "Agent 服务与 FastAPI 使用相同的 AGENT_SHARED_SECRET，并在请求头中校验 x-agent-secret。",
            "本地联调建议先启动 FastAPI，再启动 Agent 服务，最后从 Web 页面创建一场面试验证端到端链路。",
            "生产环境建议将 Agent 服务放在内网或受控网络中，仅允许 FastAPI 访问。",
            "Agent 日志应记录 session_id、状态流转、错误原因和耗时，但不应记录完整敏感简历原文。",
        ],
    )
    doc.add_heading("7.4 验证建议", level=2)
    add_bullets(
        doc,
        [
            "执行 npm run build 验证前端类型和构建。",
            "执行后端 pytest 覆盖薪酬导入和薪酬流水线测试。",
            "启动 FastAPI 后验证 GET /api/health。",
            "上传一份测试简历并确认 resume_uploads 阶段流转和候选人详情展示。",
            "创建一场面试并完成至少一轮 AI 面试，确认报告写入。",
        ],
    )
    add_design_full_appendix(doc)
    add_design_extra_detail(doc)

    path = OUT_DIR / "RecruitPro智能招聘管理系统概要设计说明书.docx"
    normalize_body_alignment(doc)
    doc.save(path)
    return path


if __name__ == "__main__":
    req = create_requirements_doc()
    design = create_design_doc()
    print(req)
    print(design)
