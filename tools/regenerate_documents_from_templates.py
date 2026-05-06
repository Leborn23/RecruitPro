from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
OUT_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATE_DIR = Path(r"D:/Program Files (x86)/jilu/xwechat_files/wxid_nztvtpqbm8qt22_9e13/msg/file/2026-04")

PROJECT_NAME = "RecruitPro智能招聘管理系统"
DOC_DATE = date(2026, 4, 27).strftime("%Y年%m月%d日")


def find_template(keyword: str) -> Path:
    for path in TEMPLATE_DIR.glob("*.docx"):
        if keyword in path.name:
            return path
    raise FileNotFoundError(keyword)


REQ_TEMPLATE = Document(str(find_template("需求")))
DESIGN_TEMPLATE = Document(str(find_template("概要")))


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def new_doc(keyword: str) -> Document:
    doc = Document(str(find_template(keyword)))
    clear_body(doc)
    normalize_styles(doc)
    return doc


def normalize_styles(doc: Document) -> None:
    for style_name in ["Normal", "List Paragraph"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.15
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.bold = True
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(5)
            style.paragraph_format.keep_with_next = False
            style.paragraph_format.page_break_before = False


def set_run_font(paragraph, size: float | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def p(doc: Document, text: str = "", style: str | None = None, align=None):
    para = doc.add_paragraph(text, style=style)
    if align is not None:
        para.alignment = align
    return para


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document, title: str) -> None:
    para = p(doc, PROJECT_NAME, style="Body Text 2")
    para.paragraph_format.space_before = Pt(115)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(para, 18, True)

    para = p(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.space_before = Pt(45)
    set_run_font(para, 22, True)

    para = p(doc, DOC_DATE, align=WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.space_before = Pt(150)
    set_run_font(para, 12, False)
    page_break(doc)


def append_template_table(doc: Document, source_doc: Document, index: int) -> Table:
    body = doc._body._element
    tbl = deepcopy(source_doc.tables[index]._tbl)
    sect = None
    if len(body) and body[-1].tag.endswith("sectPr"):
        sect = body[-1]
        body.remove(sect)
    body.append(tbl)
    if sect is not None:
        body.append(sect)
    return Table(tbl, doc)


def clear_table_text(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""


def fill_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = text
    for para in cell.paragraphs:
        set_run_font(para, 9, bold)


def add_revision_table(doc: Document, source_doc: Document) -> None:
    para = p(doc, "修订记录", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(para, 16, True)
    table = append_template_table(doc, source_doc, 0)
    clear_table_text(table)
    rows = [
        ["时间", "作者", "主要修订内容"],
        [DOC_DATE, "Codex", "V1.0 初版：根据 RecruitPro 项目源码、数据库迁移、运行文档和参考模板重新生成。"],
        ["", "", ""],
        ["", "", ""],
        ["", "", ""],
        ["", "", ""],
        ["", "", ""],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(table.cell(r, c), value, r == 0)
    page_break(doc)


def add_toc(doc: Document, rows: list[tuple[int, str, int]]) -> None:
    title = p(doc, "目 录", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title, 18, True)
    for level, text, page in rows:
        style = f"toc {min(level, 4)}"
        para = p(doc, f"{text}\t{page}", style=style if style in doc.styles else None)
        set_run_font(para, 10.5 if level <= 2 else 10, level == 1)
    page_break(doc)


def add_terms_table(doc: Document) -> None:
    table = append_template_table(doc, REQ_TEMPLATE, 1)
    clear_table_text(table)
    rows = [
        ["序号", "术语名称", "术语定义"],
        ["1", "RecruitPro", "智能招聘管理系统，覆盖岗位、简历、候选人、面试、薪酬和权限等招聘业务。"],
        ["2", "RLS", "Row Level Security，数据库行级安全策略，用于限制用户只能访问授权数据。"],
        ["3", "LLM", "大语言模型，用于简历结构化、岗位匹配、面试问答和评分报告生成。"],
        ["4", "OCR", "光学字符识别，用于在简历无法直接提取文本时进行备用解析。"],
        ["5", "Agent", "外部 AI 面试智能体服务，负责生成问题、追问、状态快照和最终报告。"],
        ["6", "匹配分", "候选人与岗位要求之间的综合评分，包含技能、项目、经验、学历和证据维度。"],
        ["7", "人工复核", "招聘人员对 AI 结果进行确认、修正或最终决策的业务动作。"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(table.cell(r, c), value, r == 0)
    p(doc)


def add_refs_table(doc: Document) -> None:
    table = append_template_table(doc, REQ_TEMPLATE, 2)
    clear_table_text(table)
    rows = [
        ["序号", "文档名称", "版本/日期", "备注"],
        ["1", "RecruitPro 项目源码", "2026.04", "D:/project/RecruitPro_"],
        ["2", "数据库迁移文件", "2026.04", "supabase/migrations"],
        ["3", "Agent 集成运行说明", "2026.04", "docs/agent-integration-runtime.md"],
        ["4", "AI 面试运行说明", "2026.04", "docs/ai-interview-runtime.md"],
        ["5", "权限说明", "2026.04", "docs/permissions.md"],
        ["6", "薪酬市场操作说明", "2026.04", "docs/salary-market-ops.md"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(table.cell(r, c), value, r == 0)
    p(doc)


MODULES = [
    ("登录与密码重置", "支持用户通过 Supabase Auth 登录系统，并提供忘记密码、邮件回调和重置密码流程。", "高", "用户访问登录页或重置链接。", "登录成功进入招聘指挥台；密码重置成功后可重新登录。", "登录失效、邮箱不存在、密码错误、回调参数缺失时给出明确提示。"),
    ("招聘指挥台", "展示岗位总数、进行中岗位、候选人样本、近期面试，以及岗位、候选人和面试的快速入口。", "高", "用户具备 VIEW_DASHBOARD 权限。", "用户可快速进入岗位、候选人或面试模块。", "统计接口失败时展示空状态和重试提示。"),
    ("岗位管理", "维护岗位名称、部门、地点、状态、筛选阈值、技术要求、年龄、学历和经验要求。", "高", "用户具备 MANAGE_POSITIONS 权限。", "岗位写入 active_positions，并作为筛选、匹配和面试安排基础。", "岗位字段缺失、权限不足、关联数据存在时需阻止危险操作。"),
    ("简历筛选", "批量上传 PDF/DOCX 简历，执行文本提取、结构化解析、岗位匹配、匹配分计算和人工复核。", "高", "用户选择目标岗位并上传合法简历文件。", "生成上传记录、结构化简历、项目经历、匹配结果和候选人档案。", "文件超限、OCR/LLM 失败、数据库写入失败时保留可重试状态。"),
    ("候选人管理与详情", "支持候选人列表、详情查看、匹配分解、证据片段、项目经历、风险提示、人工通过/待定/淘汰和删除。", "高", "候选人已由筛选流程生成或系统已有候选人数据。", "招聘人员可完成候选人复核和推进决策。", "候选人不存在、关联岗位缺失或权限不足时给出提示。"),
    ("面试管理与 AI 面试", "支持面试安排、房间入口、房间口令、AI 面试准备、会话、轮次、评分报告和人工确认。", "高", "候选人、岗位和面试安排存在。", "生成面试会话、问答轮次、AI 报告和人工确认结果。", "房间口令错误、Agent 不可用或报告未生成时进入可恢复状态。"),
    ("AI 面试 Agent", "通过 FastAPI 调用外部 Agent 服务，完成问题计划、追问、状态快照和最终评估报告。", "高", "AGENT_BASE_URL、AGENT_SHARED_SECRET 和 AGENT_TIMEOUT_MS 已配置。", "Agent 结果映射到 interview_sessions、interview_turns 和 interview_reports。", "Agent 超时、返回非 JSON、密钥错误或缺少 final_report 时生成待人工复核报告。"),
    ("系统设置与权限管理", "维护公司设置、AI 策略、默认模型、用户角色、权限键和超级管理员认领。", "高", "用户具备 MANAGE_SETTINGS 或超级管理员角色。", "权限和配置变更影响后续页面访问和业务流程。", "普通用户不得直接修改敏感表或服务密钥配置。"),
    ("薪酬市场与决策支持", "导入薪酬市场数据，生成标准化记录、薪酬基准和候选人薪酬画像。", "中", "管理员准备 CSV/JSON/JSONL 薪酬数据，并配置 Supabase 服务密钥。", "数据进入 market_salary_raw_records、normalized_records 和 benchmarks。", "标准化失败不丢失原始记录，样本不足时标记风险。"),
]


def add_module_table(doc: Document, module: tuple[str, str, str, str, str, str]) -> None:
    name, desc, priority, pre, post, err = module
    table = append_template_table(doc, REQ_TEMPLATE, 3)
    clear_table_text(table)
    rows = [
        ["模块描述", name],
        ["功能描述", desc],
        ["优先级", priority],
        ["输入/前置条件", pre],
        ["输出/后置条件", post],
        ["异常处理", err],
        ["权限要求", "页面入口、接口调用和数据库 RLS 共同约束；敏感操作由 FastAPI 或受控 RPC 完成。"],
        ["数据要求", "关键业务数据必须落库，AI 结果需保留证据、分数、状态快照和人工复核字段。"],
        ["验收标准", "正常流程、异常流程、权限不足流程和刷新恢复流程均能稳定执行。"],
        ["备注", "该模块说明按模板的“需求说明”表格格式编写。"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(table.cell(r, c), value, c == 0)
    p(doc)


def add_image_if_exists(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(4.8))
    cap = doc.add_paragraph(caption, style="Caption" if "Caption" in doc.styles else None)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


UI = ROOT / "ui-screenshots"
UI_IMAGES = [
    UI / "auth-login.png",
    UI / "dashboard.png",
    UI / "positions.png",
    UI / "screening.png",
    UI / "candidates.png",
    UI / "interviews.png",
    UI / "interview-room.png",
    UI / "settings.png",
    UI / "salary.png",
]


def add_requirement_doc() -> Path:
    doc = new_doc("需求")
    add_cover(doc, "需求规格说明书")
    add_revision_table(doc, REQ_TEMPLATE)
    add_toc(
        doc,
        [
            (1, "第一章 引言", 1),
            (2, "1.1 编写目的", 1),
            (2, "1.2 文档范围", 1),
            (2, "1.3 项目概要", 1),
            (2, "1.4 术语和缩写", 2),
            (2, "1.5 参考资料", 3),
            (1, "第二章 任务概述", 4),
            (2, "2.1 目标", 4),
            (2, "2.2 用户的特点", 5),
            (2, "2.3 假定和约束", 5),
            (1, "第三章 需求定义", 6),
            (2, "3.1 功能需求", 6),
            (3, "3.1.1 登录与密码重置", 6),
            (3, "3.1.2 招聘指挥台", 8),
            (3, "3.1.3 岗位管理", 10),
            (3, "3.1.4 简历筛选", 12),
            (3, "3.1.5 候选人管理与详情", 15),
            (3, "3.1.6 面试管理与 AI 面试", 18),
            (3, "3.1.7 AI 面试 Agent", 21),
            (3, "3.1.8 系统设置与权限管理", 24),
            (3, "3.1.9 薪酬市场与决策支持", 26),
            (2, "3.2 性能需求", 28),
            (2, "3.3 输入输出需求", 29),
            (2, "3.4 数据管理能力需求", 31),
            (2, "3.5 故障处理需求", 33),
            (1, "第四章 运行环境", 35),
            (1, "第五章 验收与交付需求", 37),
        ],
    )

    doc.add_heading("第一章 引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    p(doc, f"本文档按照参考模板格式编写，用于明确{PROJECT_NAME}的业务目标、功能范围、非功能需求、输入输出、数据管理、故障处理和运行环境，为概要设计、详细设计、开发、测试、部署和验收提供依据。")
    doc.add_heading("1.2 文档范围", level=2)
    p(doc, "本文档覆盖 Web 前端、FastAPI 业务后端、Supabase 数据与存储、AI 简历筛选、AI 面试 Agent、薪酬市场数据处理、权限与系统设置等功能。")
    doc.add_heading("1.3 项目概要", level=2)
    p(doc, f"待开发的软件系统名称：{PROJECT_NAME}。")
    p(doc, "本系统面向招聘负责人、HR、面试官、管理员和候选人，提供从岗位发布、简历上传、AI 筛选、候选人复核、面试安排、Agent 面试、报告确认到薪酬参考的招聘闭环。")
    doc.add_heading("1.4 术语和缩写", level=2)
    add_terms_table(doc)
    doc.add_heading("1.5 参考资料", level=2)
    add_refs_table(doc)

    doc.add_heading("第二章 任务概述", level=1)
    doc.add_heading("2.1 目标", level=2)
    for text in [
        "建立岗位、简历、候选人、面试、薪酬和权限统一管理的招聘业务平台。",
        "通过 AI 简历解析、岗位匹配、证据解释和人工复核提升筛选效率和透明度。",
        "通过 AI 面试 Agent 实现结构化问答、追问、评分报告和人工确认。",
        "通过 Supabase RLS、权限键和后端服务密钥控制招聘数据安全。",
        "通过薪酬市场数据和候选人薪酬画像辅助招聘 offer 决策。",
    ]:
        doc.add_paragraph(text, style="List Paragraph")
    doc.add_heading("2.2 用户的特点", level=2)
    for text in [
        "招聘负责人：关注岗位进度、候选人质量、面试安排和录用建议。",
        "HR/招聘专员：负责简历上传、筛选复核、候选人推进和面试邀约。",
        "面试官：关注候选人项目证据、面试问题、评分报告和风险点。",
        "系统管理员：负责组织设置、AI 策略、账号权限、模型配置和数据治理。",
        "候选人：通过面试房间参与在线 AI 面试，仅接触面试相关入口。",
    ]:
        doc.add_paragraph(text, style="List Paragraph")
    doc.add_heading("2.3 假定和约束", level=2)
    for text in [
        "系统需可访问 Supabase 项目、PostgreSQL、Storage、Auth、RPC 和 RLS 策略。",
        "AI 简历筛选和 AI 面试依赖可用的 LLM、OCR 和外部 Agent 服务。",
        "生产环境需配置 VITE_SUPABASE_URL、VITE_SUPABASE_ANON_KEY、SUPABASE_SERVICE_ROLE_KEY、AGENT_BASE_URL、AGENT_SHARED_SECRET 等环境变量。",
        "服务密钥只能在后端或受控脚本中使用，不得进入浏览器端。",
        "关键业务流程需要保留状态、错误码、错误信息和可重试能力。",
    ]:
        p(doc, text)

    doc.add_heading("第三章 需求定义", level=1)
    doc.add_heading("3.1 功能需求", level=2)
    for idx, module in enumerate(MODULES, start=1):
        doc.add_heading(f"3.1.{idx} {module[0]}", level=3)
        doc.add_heading("3.1.%d.1 需求说明" % idx, level=4)
        add_module_table(doc, module)
        doc.add_heading("3.1.%d.2 页面图例" % idx, level=4)
        add_image_if_exists(doc, UI_IMAGES[idx - 1], f"图 3.1.{idx}-1 {module[0]}页面")
        doc.add_heading("3.1.%d.3 详细业务规则" % idx, level=4)
        for text in [
            "进入模块前必须校验登录态和权限，后端接口不得信任前端传入的权限状态。",
            "模块核心数据必须与岗位、候选人、上传记录、面试记录或用户记录建立可追踪关联。",
            "涉及 AI 输出的结论必须保留证据和人工复核入口，不得直接作为最终录用结论。",
            "模块应支持正常流程、异常流程、刷新恢复流程和权限不足流程的验收。",
        ]:
            p(doc, text)

    doc.add_heading("3.2 性能需求", level=2)
    for text in [
        "常规列表和详情查询在正常网络和数据库负载下应在 3 秒内返回。",
        "页面交互应在 1 秒内给出加载、成功或失败反馈。",
        "批量简历处理、Agent 面试和薪酬导入属于长流程，必须具备阶段状态、错误记录和重试能力。",
        "外部 AI、OCR 和 Agent 调用必须配置超时，失败时不应造成业务数据丢失。",
    ]:
        p(doc, text)
    doc.add_heading("3.3 输入输出需求", level=2)
    for text in [
        "账号登录输入邮箱和密码，输出登录态、权限集合和错误提示。",
        "岗位管理输入岗位基础信息和筛选规则，输出岗位列表、岗位详情和结构化岗位要求。",
        "简历筛选输入简历文件和目标岗位，输出结构化简历、匹配分、证据、候选人记录和复核状态。",
        "AI 面试输入候选人上下文、岗位上下文和候选人回答，输出问题计划、追问、状态快照和评分报告。",
        "薪酬导入输入 CSV/JSON/JSONL 数据，输出原始记录、标准化记录、薪酬基准和导入任务状态。",
    ]:
        p(doc, text)
    doc.add_heading("3.4 数据管理能力需求", level=2)
    for text in [
        "系统需管理岗位、候选人、简历上传、结构化简历、岗位要求、匹配结果、面试安排、面试会话、问答轮次、面试报告、薪酬市场数据、模型配置、公司设置和用户权限。",
        "简历原文件保存至 Supabase Storage，结构化结果和处理状态保存至 PostgreSQL。",
        "关键业务表启用 RLS；后台服务使用服务密钥进行受控写入，前端按权限读取和操作。",
        "AI 输出应保留原始载荷摘要、置信度、证据片段和人工复核字段，以支持追溯。",
    ]:
        p(doc, text)
    doc.add_heading("3.5 故障处理需求", level=2)
    for text in [
        "上传、OCR、LLM、数据库写入、权限校验、外部 Agent 调用失败时需返回可读错误信息。",
        "简历处理流程需记录 pipeline_stage、status、error_code、error_message 和 retry_count。",
        "面试流程异常中断时应保留已产生的会话、轮次和报告状态，便于人工处理。",
        "Agent 未返回 final_report 时，系统需生成 pending_human_review 报告，避免结果丢失。",
    ]:
        p(doc, text)

    doc.add_heading("第四章 运行环境", level=1)
    for heading, body in [
        ("4.1 设备", "客户端为现代浏览器；服务端可部署于 Linux 或容器环境；数据库和对象存储依赖 Supabase。"),
        ("4.2 支持软件", "前端采用 React 19、TypeScript、Vite、Tailwind CSS、React Router、Supabase JS；后端采用 FastAPI、Python、Supabase Python Client、httpx。"),
        ("4.3 接口", "系统接口包括前端到 Supabase、前端到 FastAPI、FastAPI 到 Supabase、FastAPI 到 Agent/LLM/OCR 等集成接口。"),
        ("4.4 控制", "系统通过浏览器人机交互控制业务流程，通过权限键控制页面和操作可见性，通过后端接口与数据库策略控制敏感数据写入。"),
    ]:
        doc.add_heading(heading, level=2)
        p(doc, body)

    doc.add_heading("第五章 验收与交付需求", level=1)
    for title, body in [
        ("5.1 功能验收", "岗位、筛选、候选人、面试、Agent、薪酬、权限和设置模块均需完成正常、异常、权限不足和刷新恢复场景验证。"),
        ("5.2 数据验收", "验收环境需准备至少一个岗位、多份简历、多个候选人、一次完整 AI 面试、一份人工确认报告和一组薪酬样例数据。"),
        ("5.3 交付清单", "交付内容包括源码、数据库迁移、运行文档、需求规格说明书、概要设计说明书、部署说明、测试记录和遗留问题清单。"),
        ("5.4 最终确认", "项目验收完成后，应记录验收日期、验收人、验收结论和遗留事项。"),
    ]:
        doc.add_heading(title, level=2)
        p(doc, body)

    path = OUT_DIR / "RecruitPro智能招聘管理系统需求规格说明书.docx"
    doc.save(path)
    return path


def add_design_doc() -> Path:
    doc = new_doc("概要")
    add_cover(doc, "概要设计说明书")
    add_toc(
        doc,
        [
            (1, "第一章 引言", 1),
            (2, "1.1 编写目的", 1),
            (2, "1.2 文档范围", 1),
            (2, "1.3 项目概要", 1),
            (2, "1.4 术语和缩写", 2),
            (2, "1.5 参考资料", 3),
            (2, "1.6 修订记录", 4),
            (1, "第二章 总体设计", 5),
            (2, "2.1 需求规定", 5),
            (2, "2.2 运行环境", 8),
            (2, "2.3 基本设计概念和处理流程", 9),
            (2, "2.4 系统架构", 11),
            (2, "2.5 功能模块设计", 12),
            (1, "第三章 接口设计", 22),
            (1, "第四章 运行设计", 27),
            (1, "第五章 系统数据结构设计", 32),
            (1, "第六章 系统安全设计", 38),
            (1, "第七章 部署与运维设计", 42),
        ],
    )

    doc.add_heading("第一章 引言", level=1)
    for heading, body in [
        ("1.1 编写目的", f"本文档按照参考模板格式编写，在需求规格说明基础上，对{PROJECT_NAME}的总体架构、模块职责、接口、运行流程、数据结构、安全和部署方式进行概要设计说明。"),
        ("1.2 文档范围", "本文档覆盖前端 React 应用、FastAPI 后端、Supabase 数据层、AI 简历筛选、AI 面试 Agent、薪酬市场数据处理、权限和系统设置。"),
        ("1.3 项目概要", "RecruitPro 是面向招聘团队的智能招聘管理系统，围绕岗位、简历、候选人、面试、薪酬和权限形成闭环。系统采用前后端分离架构，复杂业务和敏感写入集中在 FastAPI 后端。"),
    ]:
        doc.add_heading(heading, level=2)
        p(doc, body)
    doc.add_heading("1.4 术语和缩写", level=2)
    add_terms_table(doc)
    doc.add_heading("1.5 参考资料", level=2)
    add_refs_table(doc)
    doc.add_heading("1.6 修订记录", level=2)
    table = append_template_table(doc, REQ_TEMPLATE, 0)
    clear_table_text(table)
    for r, row in enumerate([["时间", "作者", "主要修订内容"], [DOC_DATE, "Codex", "按概要设计模板格式重新生成。"], ["", "", ""], ["", "", ""], ["", "", ""], ["", "", ""], ["", "", ""]]):
        for c, value in enumerate(row):
            fill_cell(table.cell(r, c), value, r == 0)
    p(doc)

    doc.add_heading("第二章 总体设计", level=1)
    doc.add_heading("2.1 需求规定", level=2)
    for idx, module in enumerate(MODULES, start=1):
        doc.add_heading(f"2.1.{idx} {module[0]}", level=3)
        doc.add_heading(f"2.1.{idx}.1 需求说明", level=4)
        add_module_table(doc, module)
        doc.add_heading(f"2.1.{idx}.2 页面图例", level=4)
        add_image_if_exists(doc, UI_IMAGES[idx - 1], f"图 2.1.{idx}-1 {module[0]}页面")
    doc.add_heading("2.2 运行环境", level=2)
    p(doc, "浏览器端运行 React 单页应用；后端运行 FastAPI；数据库采用 Supabase PostgreSQL、Auth、Storage、RPC 和 RLS；外部 AI 服务包括 LLM、OCR 和 Agent。")
    doc.add_heading("2.3 基本设计概念和处理流程", level=2)
    for text in [
        "以岗位为业务中心，岗位规则驱动简历筛选、候选人匹配和面试推进。",
        "以证据为 AI 输出基础，AI 解析和评分需要保留证据、结构化字段、置信度和人工复核状态。",
        "以前后端分离为工程边界，展示和交互在前端，长流程和敏感操作在后端。",
        "以数据库安全为底线，权限不只依赖前端路由，还需通过 RLS、RPC 和后端接口共同约束。",
    ]:
        p(doc, text)
    doc.add_heading("2.4 系统架构", level=2)
    for text in [
        "表现层：React SPA、React Router、Tailwind CSS，负责页面路由、权限保护、数据展示、文件上传和用户交互。",
        "业务层：FastAPI，负责简历筛选、薪酬导入、面试运行、管理员接口和外部 Agent 调用。",
        "数据层：Supabase PostgreSQL，负责岗位、候选人、简历、匹配、面试、薪酬、权限和配置等结构化数据。",
        "存储层：Supabase Storage，负责简历文件和头像等对象存储。",
        "AI 服务层：LLM/OCR/Agent，负责简历解析、岗位匹配、面试问答和评分报告生成。",
    ]:
        p(doc, text)
    doc.add_heading("2.5 功能模块设计", level=2)
    for module in MODULES:
        doc.add_heading(module[0], level=3)
        for text in [
            f"前端设计：围绕{module[0]}页面入口、列表、详情、状态反馈和操作按钮组织组件。",
            "后端设计：复杂业务、服务密钥操作、长流程编排和外部服务调用集中在 FastAPI。",
            "数据设计：模块数据通过 Supabase 表结构保存，关键 AI 输出保留结构化字段和证据。",
            "安全设计：页面级权限、接口级权限和数据库 RLS 共同约束。",
        ]:
            p(doc, text)

    doc.add_heading("第三章 接口设计", level=1)
    for text in [
        "前端到 FastAPI：负责岗位、候选人、上传、面试、Agent、薪酬和管理员等业务接口。",
        "FastAPI 到 Supabase：使用 service role key 执行受控写入和复杂聚合。",
        "FastAPI 到 Agent：通过 AGENT_BASE_URL 调用 /agent/start、/agent/answer 和 /agent/status，并携带 x-agent-secret。",
        "前端到 Supabase：仅使用匿名密钥和 RLS 保护下的数据访问。",
    ]:
        p(doc, text)
    doc.add_heading("第四章 运行设计", level=1)
    for text in [
        "简历筛选流程：上传文件、文本提取、结构化解析、岗位匹配、候选人入库、人工复核。",
        "AI 面试流程：准备面试、启动 Agent、保存首问、候选人答题、Agent 追问或下一题、结束评分、人工确认。",
        "薪酬导入流程：读取原始数据、保存 raw_records、标准化 normalized_records、刷新 benchmarks。",
        "权限管理流程：超级管理员通过受控接口修改角色和权限键，前端和后端同步生效。",
    ]:
        p(doc, text)
    doc.add_heading("第五章 系统数据结构设计", level=1)
    for text in [
        "岗位域：active_positions、parsed_job_requirements。",
        "候选人域：candidates、parsed_resume_profiles、parsed_resume_projects、candidate_position_matches。",
        "面试域：upcoming_interviews、interview_sessions、interview_turns、interview_reports。",
        "薪酬域：market_salary_raw_records、market_salary_normalized_records、market_salary_benchmarks、market_salary_crawl_jobs。",
        "权限与配置域：user_roles、company_settings、llm_model_configs、llm_usage_events。",
    ]:
        p(doc, text)
    doc.add_heading("第六章 系统安全设计", level=1)
    for text in [
        "认证采用 Supabase Auth，前端页面通过 ProtectedRoute 控制访问。",
        "授权采用权限键和角色层级，高权限动作通过后端接口或受控 RPC 完成。",
        "服务密钥只允许后端和离线脚本使用，不得暴露在浏览器端。",
        "Agent 服务不直接操作业务数据库，所有持久化由 FastAPI 完成。",
        "AI 报告必须支持人工确认，最终录用建议由人工确认结果决定。",
    ]:
        p(doc, text)
    doc.add_heading("第七章 部署与运维设计", level=1)
    for text in [
        "前端执行 npm run build 后部署 dist 静态资源。",
        "FastAPI 可直接运行或容器化部署，生产环境需配置 Supabase 和 Agent 相关环境变量。",
        "数据库结构以 supabase/migrations 为来源，通过迁移同步。",
        "Agent 服务部署为 FastAPI 可访问的 HTTP 服务，并校验 x-agent-secret。",
        "发布前需完成前端构建、后端健康检查、数据库迁移检查、Agent 连通性检查和端到端招聘流程验证。",
    ]:
        p(doc, text)

    path = OUT_DIR / "RecruitPro智能招聘管理系统概要设计说明书.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(add_requirement_doc())
    print(add_design_doc())
